#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# SPDX-License-Identifier: MIT
"""
core.py — Kernlogik des anonymizer-Moduls
==========================================

PRIVAT — klientenbezogene Daten, nicht veröffentlichen.

Extrahiert aus BACH hub/_services/document/anonymizer_service.py v1.2.0 (MIT).
Alle BACH-Laufzeitbezüge neutralisiert; keine bach_paths-, bach.db- oder
BACH-Hub-Imports. Vollständig standalone lauffähig.

Herkunfts-Komponentenreferenzen (NICHT zur Laufzeit importiert — inline implementiert):
  - RedactionDetector-Logik (core/redaction/detector.py)
  - OCREngine-Aufruf (core/ocr/engine.py)
  - PDFProcessor-Pipeline (core/pdf/processor.py)

Pseudonymisierung und De-Anonymisierung von Dokumenten.
Ersetzt personenbezogene Daten durch konsistente Tarnnamen.
Schlüssel wird mit Fernet authentifiziert verschlüsselt gespeichert.

Pfad-Auflösung (neutral, kein BACH):
  Schlüsselablage : ANONYMIZER_KEYS_DIR (ENV)
                    → %LOCALAPPDATA%\\anonymizer\\keys (Windows)
                    → ~/.local/share/anonymizer/keys (Unix)
  Klienten-Ausgabe: ANONYMIZER_HOME (ENV) → ~/.anonymizer
  Whitelist       : ANONYMIZER_HOME/anonymizer_whitelist.json (optional)

Kernabhängigkeit:
  cryptography   — authentifizierte Fernet-Schlüsselverschlüsselung

Optionale Abhängigkeiten (pip install ...):
  python-docx    — Word-Dokument-Anonymisierung
  PyMuPDF        — PDF-Schwärzung (fitz)
  pikepdf        — PDF-AES-256-Verschlüsselung
  openpyxl       — Excel-Anonymisierung

Version: 0.2.1 (defusedxml-Härtung 2026-07-23; Basis: 0.2.0 Sicherheits- und
Datenschutz-Härtung 2026-07-16)
Copyright (c) 2026 ellmos / BACH Contributors — MIT License
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import secrets
import shutil
import stat
import tempfile
import zipfile
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, FrozenSet, List, Optional, Tuple
from defusedxml import ElementTree


# Authentifizierte Schlüsselverschlüsselung
try:
    from cryptography.fernet import Fernet
    from cryptography.hazmat.primitives import hashes
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    import base64
    CRYPTO_AVAILABLE = True
except ImportError:
    CRYPTO_AVAILABLE = False

# Word-Dokumente
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF-Verarbeitung (PyMuPDF)
try:
    import fitz
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

# PDF-Verschlüsselung (pikepdf)
try:
    import pikepdf
    PIKEPDF_AVAILABLE = True
except ImportError:
    PIKEPDF_AVAILABLE = False

# Excel-Verarbeitung (openpyxl)
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Personennamen-Erkennung (spaCy NER, DE+EN) -- generische Erkennung von
# Personennamen im Fliesstext, unabhaengig von Schreibweise/Sprache/Herkunft.
# Robuster als Regex-Heuristiken, da PERSON von Fachbegriffen/Orten/Institutionen
# unterschieden wird (im Deutschen wird sonst JEDES Substantiv grossgeschrieben).
#
# WICHTIG: spaCy importiert transitiv `cProfile` -> `import profile` (Stdlib).
# BACH hat ein EIGENES Modul `hub/profile.py` -- landet dessen Ordner (statt
# nur `system/`) auf sys.path (z.B. durch Test-Sammlung anderer Dateien),
# verschattet es die Stdlib und `cProfile` bricht mit AttributeError ab
# (nicht ImportError!). Fix: das ECHTE Stdlib-`profile`-Modul vorab explizit
# laden und in sys.modules cachen, BEVOR spaCy/cProfile es importieren --
# das gewinnt garantiert gegen jede sys.path-Reihenfolge.
try:
    import sys as _sys
    if "profile" not in _sys.modules:
        import importlib.util as _importlib_util
        import sysconfig as _sysconfig
        _stdlib_profile_path = os.path.join(_sysconfig.get_path("stdlib"), "profile.py")
        if os.path.exists(_stdlib_profile_path):
            _spec = _importlib_util.spec_from_file_location("profile", _stdlib_profile_path)
            _stdlib_profile_mod = _importlib_util.module_from_spec(_spec)
            _spec.loader.exec_module(_stdlib_profile_mod)
            _sys.modules["profile"] = _stdlib_profile_mod

    import spacy
    SPACY_AVAILABLE = True
except Exception as _spacy_import_error:
    SPACY_AVAILABLE = False
    print(f"[WARN] spaCy nicht verfuegbar ({_spacy_import_error}) -- NER-Namenserkennung deaktiviert")


# ═══════════════════════════════════════════════════════════════
# Datenklassen
# ═══════════════════════════════════════════════════════════════

@dataclass
class AnonymProfile:
    """Anonymisierungsprofil für einen Klienten."""
    client_id: str                          # z.B. "K_0042"
    tarnname: str                           # z.B. "Felix Bergmann"
    fake_geburtsdatum: str                  # z.B. "22.07.2016" (gleiches Alter)
    mappings: Dict[str, Dict[str, str]] = field(default_factory=dict)
    # mappings = {
    #   "names": {"Max Mustermann": "Felix Bergmann", "Dr. Meyer": "Dr. Lindner"},
    #   "dates": {"15.03.2016": "22.07.2016"},
    #   "addresses": {"Musterstr. 5": "Waldweg 12"},
    #   "misc": {"07761/123456": "07741/987654"}
    # }
    created: str = ""
    version: int = 1


@dataclass
class AnonymResult:
    """Ergebnis einer Anonymisierung."""
    processed_files: int = 0
    anonymized_files: int = 0
    skipped_files: int = 0
    errors: List[str] = field(default_factory=list)
    replacements_total: int = 0


@dataclass
class ProgressInfo:
    """Fortschrittsinfo für GUI."""
    total_files: int = 0
    processed_files: int = 0
    current_file: str = ""
    status: str = "idle"  # idle, scanning, anonymizing, done, error

    @property
    def percent(self) -> float:
        if self.total_files == 0:
            return 0.0
        return (self.processed_files / self.total_files) * 100.0


# ═══════════════════════════════════════════════════════════════
# Tarnnamen-Generator
# ═══════════════════════════════════════════════════════════════

# Deutsche Phantasienamen für konsistente Pseudonymisierung
# Getrennt nach Geschlecht für konsistente Pronomen
TARN_VORNAMEN_M = [
    "Felix", "Jonas", "Paul", "Leon", "Finn", "Lukas", "Noah", "Elias",
    "Ben", "Maximilian", "Tim", "David", "Moritz", "Julian", "Niklas", "Erik"
]

TARN_VORNAMEN_W = [
    "Lena", "Marie", "Sophie", "Emma", "Hannah", "Mia", "Clara", "Lina",
    "Lea", "Anna", "Laura", "Sarah", "Julia", "Amara", "Emily", "Nina"
]

# Kombinierte Liste für Rückwärtskompatibilität
TARN_VORNAMEN = TARN_VORNAMEN_M + TARN_VORNAMEN_W

# Häufige deutsche Vornamen für Gender-Erkennung
DEUTSCHE_VORNAMEN_M = {
    "jaden", "max", "paul", "leon", "jonas", "felix", "lukas", "elias", "ben", "noah",
    "tim", "david", "jan", "finn", "niklas", "moritz", "julian", "erik", "tom", "luca",
    "michael", "thomas", "peter", "hans", "klaus", "christian", "andreas", "stefan",
    "markus", "daniel", "tobias", "sebastian", "florian", "matthias", "alexander",
    "johannes", "philipp", "simon", "marco", "oliver", "martin", "frank", "wolfgang",
    "karl", "josef", "heinrich", "werner", "günter", "jürgen", "horst", "dieter",
    "harald", "helmut", "manfred", "bernhard", "gerhard", "rainer", "rolf", "walter",
    # Zusätzliche Namen für bessere Erkennung
    "skaven", "kevin", "justin", "jason", "brandon", "tyler", "ryan", "kyle",
    "pascal", "dennis", "sven", "lars", "nils", "björn", "jens", "uwe", "kai"
}

DEUTSCHE_VORNAMEN_W = {
    "marie", "sophie", "emma", "lena", "hannah", "mia", "clara", "lina", "lea", "anna",
    "laura", "sarah", "julia", "lisa", "emily", "nina", "amelie", "leonie", "johanna",
    "maria", "sabine", "petra", "susanne", "monika", "karin", "ursula", "renate",
    "brigitte", "helga", "ingrid", "gisela", "erika", "christa", "hildegard", "gertrud",
    "elisabeth", "christine", "andrea", "claudia", "martina", "nicole", "katrin",
    "birgit", "silke", "heike", "anja", "melanie", "stefanie", "sandra", "jennifer"
}

TARN_NACHNAMEN = [
    "Bergmann", "Fischer", "Lindner", "Sommer", "Richter", "Vogel",
    "Baumann", "Krause", "Werner", "Hartmann", "Lehmann", "Brandt",
    "Keller", "Bauer", "Schuster", "Hofmann", "Albrecht", "Steiner"
]

TARN_STRASSEN = [
    "Waldweg", "Birkenallee", "Sonnenstr.", "Gartenweg", "Lindenstr.",
    "Bergstr.", "Rosenweg", "Eichenstr.", "Parkstr.", "Wiesenweg"
]

TARN_STAEDTE = [
    "79800 Tiengen", "79761 Waldshut", "79725 Laufenburg",
    "79780 Stühlingen", "79807 Lottstetten", "79737 Herrischried"
]


def _detect_gender(vorname: str) -> str:
    """
    Erkennt das Geschlecht anhand des Vornamens.

    Returns:
        'm' für männlich, 'w' für weiblich, 'u' für unbekannt
    """
    vorname_lower = vorname.lower().strip()

    if vorname_lower in DEUTSCHE_VORNAMEN_M:
        return 'm'
    if vorname_lower in DEUTSCHE_VORNAMEN_W:
        return 'w'

    # Heuristiken für unbekannte Namen
    # Viele weibliche Namen enden auf -a, -e, -ie, -ine
    if vorname_lower.endswith(('a', 'ie', 'ine', 'ette', 'ella', 'ina')):
        return 'w'
    # Viele männliche Namen enden auf Konsonanten oder -o, -us, -er
    if vorname_lower.endswith(('us', 'er', 'o', 'ian', 'en')):
        return 'm'

    return 'u'


def _generate_tarnname(used_names: set = None, gender: str = None, original_vorname: str = None) -> str:
    """
    Generiert einen zufälligen Tarnnamen.

    Args:
        used_names: Set bereits verwendeter Namen (Kollisionsvermeidung)
        gender: 'm' für männlich, 'w' für weiblich, None für auto-detect
        original_vorname: Echter Vorname für Gender-Erkennung

    Returns:
        Tarnname im Format "Vorname Nachname"
    """
    if used_names is None:
        used_names = set()

    # Gender automatisch erkennen wenn nicht angegeben
    if gender is None and original_vorname:
        gender = _detect_gender(original_vorname)

    # Passende Vornamenliste wählen
    if gender == 'm':
        vornamen_pool = TARN_VORNAMEN_M
    elif gender == 'w':
        vornamen_pool = TARN_VORNAMEN_W
    else:
        vornamen_pool = TARN_VORNAMEN  # Fallback: alle

    for _ in range(100):
        name = f"{secrets.choice(vornamen_pool)} {secrets.choice(TARN_NACHNAMEN)}"
        if name not in used_names:
            return name
    return f"Person_{secrets.token_hex(4)}"


def _shift_date(date_str: str, days_offset: int) -> str:
    """
    Verschiebt ein Datum um eine feste Anzahl Tage.
    Akzeptiert dd.mm.yyyy Format.
    """
    try:
        dt = datetime.strptime(date_str.strip(), "%d.%m.%Y")
        shifted = dt + timedelta(days=days_offset)
        return shifted.strftime("%d.%m.%Y")
    except ValueError:
        return date_str


def _generate_fake_date_same_age(real_date: str) -> Tuple[str, int]:
    """
    Generiert ein falsches Geburtsdatum mit demselben Alter.
    Returns: (fake_date, days_offset)
    """
    offset = secrets.randbelow(300) - 150  # -150 bis +149 Tage
    fake = _shift_date(real_date, offset)
    return fake, offset


def _generate_fake_phone() -> str:
    """Generiert eine falsche Telefonnummer."""
    prefix = secrets.choice(["07741", "07742", "07743", "07744"])
    number = "".join([str(secrets.randbelow(10)) for _ in range(6)])
    return f"{prefix}/{number}"


def _generate_fake_email(original: str = None) -> str:
    """Generiert eine falsche E-Mail-Adresse."""
    vorname = secrets.choice(TARN_VORNAMEN).lower()
    nachname = secrets.choice(TARN_NACHNAMEN).lower()
    domain = secrets.choice(["beispiel.de", "muster.org", "test-mail.de", "privat.net"])
    return f"{vorname}.{nachname}@{domain}"


def _generate_fake_address() -> str:
    """Generiert eine falsche Adresse."""
    strasse = secrets.choice(TARN_STRASSEN)
    nr = secrets.randbelow(50) + 1
    stadt = secrets.choice(TARN_STAEDTE)
    return f"{strasse} {nr}, {stadt}"


_TARN_INSTITUTION_PREFIXES = [
    "Linden", "Eichen", "Birken", "Ahorn", "Buchen", "Tannen",
    "Sonnen", "Mond", "Stern", "Regen", "Bach", "Berg", "Wald",
]

def _generate_fake_institution(original: str) -> str:
    """Generiert einen falschen Institutionsnamen mit gleichem Typ-Suffix."""
    # Suffix extrahieren (schule, klinik, heim, etc.)
    suffixes = ["schule", "klinik", "heim", "werkstatt", "zentrum", "praxis",
                "kindergarten", "kita", "hort", "internat", "wohnheim",
                "wohngruppe", "tagesstaette", "foerderzentrum"]
    found_suffix = ""
    for s in suffixes:
        if s in original.lower():
            found_suffix = s
            break
    if not found_suffix:
        found_suffix = "schule"
    prefix = secrets.choice(_TARN_INSTITUTION_PREFIXES)
    stadt = secrets.choice(TARN_STAEDTE)
    return f"{prefix}{found_suffix} {stadt}"


def _replace_word_boundary(text: str, old: str, new: str) -> Tuple[str, int]:
    """
    Ersetzt `old` in `text` NUR an Wortgrenzen (nicht als blinder Teilstring).

    Verhindert kaputte Fragmente wie "Person026sbesonderheiten" statt
    "Wahrnehmungsbesonderheiten": Ein NER-erkannter Einzelwort-Name kann
    zufaellig Praefix eines laengeren deutschen Kompositums sein -- an DER
    Fundstelle greift die Wortgrenzen-Pruefung bei der Erkennung, aber die
    anschliessende Ersetzung ist global (jedes Vorkommen im Dokument), auch
    an Stellen, wo derselbe Teilstring OHNE Wortgrenze auftaucht.

    Python-`\\b` behandelt deutsche Umlaute/ß korrekt als Wortzeichen.
    """
    if not old:
        return text, 0
    # RFC email domains are case-insensitive and real documents commonly vary
    # the local-part casing as well. Keep name replacement conservative, but
    # make discovered email mappings cover every spelling variant.
    flags = re.IGNORECASE if EMAIL_PATTERN.fullmatch(old) else 0
    pattern = re.compile(r'\b' + re.escape(old) + r'\b', flags)
    count = len(pattern.findall(text))
    if count:
        text = pattern.sub(lambda m: new, text)
    return text, count


# ═══════════════════════════════════════════════════════════════
# RegEx-Patterns für automatische Erkennung
# ═══════════════════════════════════════════════════════════════

# Telefonnummern (deutsch): 07761/123456, +49 7761 123456, 0049-7761-123456, 0761 12345678
PHONE_PATTERN = re.compile(
    r'(?:'
    r'(?:\+49|0049)[\s\-/]?\d{2,4}[\s\-/]?\d{4,8}|'  # +49 oder 0049
    r'0\d{2,4}[\s\-/]?\d{4,8}'                        # 0xxx/xxxxxxx
    r')',
    re.IGNORECASE
)

# E-Mail-Adressen
EMAIL_PATTERN = re.compile(
    r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
    re.IGNORECASE
)

# Straßenadressen: "Musterstr. 5", "Hauptstraße 123a", "Am Waldweg 7 b"
STREET_PATTERN = re.compile(
    r'(?:'
    r'(?:[A-ZÄÖÜ][a-zäöüß]+(?:str\.|straße|weg|gasse|platz|allee|ring|damm|ufer|berg|tal|hof|feld|wiese|grund|rain|steig|pfad))'  # Straßenname
    r'\s*'
    r'\d{1,4}\s?[a-zA-Z]?'  # Hausnummer + opt. Zusatz
    r')',
    re.IGNORECASE
)

# Institutionsnamen (Schulen, Kliniken, Einrichtungen) + Ortsnamen
INSTITUTION_PATTERN = re.compile(
    r'(?:[A-ZÄÖÜ][a-zäöüß]+(?:schule|klinik|heim|werkstatt|zentrum|praxis|kindergarten|kita|hort|internat|wohnheim|wohngruppe|tagesstaette|foerderzentrum))'
    r'(?:\s+[A-ZÄÖÜ][a-zäöüß]+(?:\s+[a-zäöüß]+)?)?',  # Optionaler Ortsname
    re.UNICODE
)

# Personennamen in Tabellenzeilen (Teilnehmerlisten, Gruppen-Protokolle):
# Faengt Zeilen wie "Timon Ackerknecht | Teilt seine technische Expertise..."
# oder "Amara Wanjiru Osei Boateng | Sehr fixiert auf Essen...".
# Eng verankert (Zeilenanfang + 2-4 grossgeschriebene Woerter + unmittelbar
# gefolgt von einem Tabellen-Trennzeichen), um Falsch-Positive bei normaler
# (grossschreibungsreicher) deutscher Fliesstext-Prosa zu vermeiden.
# Erfasst KEINE explizit uebergebenen Namen (Klient/Eltern) -- das ist die
# einzige automatische Erkennung fuer unbekannte Drittpersonen (z.B. andere
# Kinder in Gruppenprotokollen), die sonst nirgends erfasst werden.
TABLE_ROW_NAME_PATTERN = re.compile(
    r'^([A-ZÄÖÜ][a-zäöüß]+(?:[\s\-][A-ZÄÖÜ][a-zäöüß]+){1,3})\s*[|\t]',
    re.MULTILINE
)

# Historische Domain-Liste. Die Erkennung filtert seit 0.2.0 nicht mehr danach:
# jede syntaktisch gueltige Adresse ist im Datenschutzkontext sensibel.
_PRIVATE_EMAIL_DOMAINS = {
    "gmail.com", "gmx.de", "gmx.net", "web.de", "yahoo.de", "yahoo.com",
    "hotmail.com", "hotmail.de", "outlook.com", "outlook.de", "live.de",
    "live.com", "t-online.de", "freenet.de", "arcor.de", "aol.com",
    "icloud.com", "me.com", "mail.de", "email.de", "posteo.de",
    "kabelbw.de", "kabelmail.de", "vodafone.de", "o2online.de",
    "unitymedia.de", "1und1.de", "ionos.de",
}

SUPPORTED_SUFFIXES = frozenset({".docx", ".txt", ".md", ".pdf", ".xlsx", ".doc"})
_CLIENT_ID_PATTERN = re.compile(r"K_[A-Za-z0-9][A-Za-z0-9_-]{2,63}\Z")
_MAX_INPUT_BYTES = 100 * 1024 * 1024
_MAX_KEY_FILE_BYTES = 10 * 1024 * 1024
_MAX_ARCHIVE_ENTRIES = 10_000
_MAX_ARCHIVE_UNCOMPRESSED_BYTES = 500 * 1024 * 1024
_MAX_ARCHIVE_RATIO = 200
_CLOUD_PATH_MARKERS = (
    "onedrive", "dropbox", "google drive", "googledrive", "icloud drive",
    "box sync", "nextcloud",
)
_UNVERIFIED_PACKAGE_PREFIXES = (
    "word/media/", "word/embeddings/", "word/activeX/",
    "xl/media/", "xl/embeddings/", "xl/activeX/",
)
# Nur Bild-/Medien-Einträge (nicht Embeddings/ActiveX) dürfen ueberhaupt gegen
# ein vertrauenswürdiges Template freigegeben werden; die anderen Praefixe in
# _UNVERIFIED_PACKAGE_PREFIXES bleiben unconditional gesperrt.
_TRUSTED_MEDIA_PREFIXES = ("word/media/", "xl/media/")


def _is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
        return True
    except ValueError:
        return False


def _is_cloud_path(path: Path) -> bool:
    """Conservative local-path guard for secrets and restored plaintext."""
    resolved = str(path.expanduser().resolve(strict=False)).casefold()
    if any(marker in resolved for marker in _CLOUD_PATH_MARKERS):
        return True
    for env_name in ("OneDrive", "OneDriveConsumer", "OneDriveCommercial"):
        root = os.environ.get(env_name)
        if root and _is_relative_to(
            path.expanduser().resolve(strict=False),
            Path(root).expanduser().resolve(strict=False),
        ):
            return True
    return False


def _is_link_like(path: Path) -> bool:
    """Return true for symlinks and Windows junction/reparse-point objects."""
    try:
        metadata = os.lstat(path)
    except OSError:
        return False
    if stat.S_ISLNK(metadata.st_mode):
        return True
    attributes = getattr(metadata, "st_file_attributes", 0)
    reparse_flag = getattr(stat, "FILE_ATTRIBUTE_REPARSE_POINT", 0x400)
    return bool(attributes & reparse_flag)


def _reject_symlink(path: Path, *, allow_missing_leaf: bool = False) -> None:
    """Reject link-like components so checks and writes address one object."""
    candidate = path.expanduser().absolute()
    parts = candidate.parts
    current = Path(parts[0])
    for index, part in enumerate(parts[1:], start=1):
        current = current / part
        if not current.exists():
            if allow_missing_leaf and index == len(parts) - 1:
                return
            continue
        if _is_link_like(current):
            raise ValueError("Links and reparse points are not accepted at this boundary")


def _validate_client_id(client_id: str) -> str:
    if not isinstance(client_id, str) or not _CLIENT_ID_PATTERN.fullmatch(client_id):
        raise ValueError("Invalid client_id; expected K_ followed by 3-64 safe characters")
    return client_id


def _validate_regular_input(path: Path) -> Path:
    _reject_symlink(path)
    if not path.is_file():
        raise ValueError("Input must be an existing regular file")
    size = path.stat().st_size
    if size > _MAX_INPUT_BYTES:
        raise ValueError("Input exceeds the 100 MiB safety limit")
    return path.resolve(strict=True)


def _collect_replacements(profile: AnonymProfile) -> List[Tuple[str, str]]:
    merged: Dict[str, str] = {}
    for category in profile.mappings.values():
        if isinstance(category, dict):
            for old, new in category.items():
                if isinstance(old, str) and isinstance(new, str) and old:
                    merged[old] = new
    return sorted(merged.items(), key=lambda item: len(item[0]), reverse=True)


def _replace_mapped_text(text: str, replacements: List[Tuple[str, str]]) -> Tuple[str, int]:
    count = 0
    for old, new in replacements:
        text, occurrences = _replace_word_boundary(text, old, new)
        count += occurrences
    return text, count


def _residual_originals(text: str, replacements: List[Tuple[str, str]]) -> List[str]:
    residuals = []
    for old, _ in replacements:
        flags = re.IGNORECASE if EMAIL_PATTERN.fullmatch(old) else 0
        pattern = re.compile(r'\b' + re.escape(old) + r'\b', flags)
        if pattern.search(text):
            residuals.append(old)
    return residuals


def _validate_archive(path: Path) -> List[zipfile.ZipInfo]:
    with zipfile.ZipFile(path, "r") as archive:
        infos = archive.infolist()
        if len(infos) > _MAX_ARCHIVE_ENTRIES:
            raise ValueError("Archive contains too many entries")
        total = 0
        for info in infos:
            normalized = info.filename.replace("\\", "/")
            if normalized.startswith("/") or ".." in Path(normalized).parts:
                raise ValueError("Archive contains an unsafe member path")
            total += info.file_size
            if total > _MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                raise ValueError("Archive expands beyond the safety limit")
            if info.compress_size and info.file_size / info.compress_size > _MAX_ARCHIVE_RATIO:
                raise ValueError("Archive entry exceeds the compression-ratio limit")
        return infos


def _has_unverified_package_content(
    path: Path, trusted_media_hashes: Optional[FrozenSet[str]] = None
) -> bool:
    """True wenn die OOXML-Datei unverifizierten Medien-/Makroinhalt enthaelt.

    word/media|xl/media-Eintraege gelten als verifiziert, wenn ihr SHA-256
    byte-identisch in trusted_media_hashes vorkommt (siehe
    _load_trusted_template_media_hashes). Embeddings/ActiveX-Eintraege und
    VBA-/OLE-Payloads bleiben davon unberuehrt und sperren immer.
    """
    infos = _validate_archive(path)
    with zipfile.ZipFile(path, "r") as archive:
        for info in infos:
            name = info.filename.replace("\\", "/").lstrip("/")
            if name.casefold().endswith(("vbaproject.bin", ".ole")):
                return True
            if name.startswith(_TRUSTED_MEDIA_PREFIXES):
                if trusted_media_hashes and hashlib.sha256(
                    archive.read(info)
                ).hexdigest() in trusted_media_hashes:
                    continue
                return True
            if name.startswith(_UNVERIFIED_PACKAGE_PREFIXES):
                return True
    return False


def _load_trusted_template_media_hashes(
    template_path: Optional[str],
) -> Optional[FrozenSet[str]]:
    """Berechnet SHA-256-Hashes aller word/media|xl/media-Eintraege eines
    vertrauenswuerdigen Templates.

    Gibt None zurueck, wenn kein Template konfiguriert ist (unveraendertes
    Verhalten: jeder Medien-Eintrag in verarbeiteten Dateien sperrt weiterhin).
    Ein konfiguriertes, aber ungueltiges/fehlendes Template wirft bewusst
    (fail-closed statt eine vermeintlich aktive Freigabe still zu ignorieren).
    """
    if not template_path:
        return None
    path = _validate_regular_input(Path(template_path))
    infos = _validate_archive(path)
    hashes: set = set()
    with zipfile.ZipFile(path, "r") as archive:
        for info in infos:
            name = info.filename.replace("\\", "/").lstrip("/")
            if name.startswith(_TRUSTED_MEDIA_PREFIXES):
                hashes.add(hashlib.sha256(archive.read(info)).hexdigest())
    return frozenset(hashes)


def _rewrite_xml_payload(data: bytes, replacements: List[Tuple[str, str]]) -> Tuple[bytes, int]:
    try:
        root = ElementTree.fromstring(data)
    except ElementTree.ParseError:
        decoded = data.decode("utf-8", errors="strict")
        replaced, count = _replace_mapped_text(decoded, replacements)
        return replaced.encode("utf-8"), count

    count = 0
    # Word commonly divides one visible name over several runs. Collapse only
    # paragraph text nodes when a mapped value crosses those run boundaries.
    for paragraph in [elem for elem in root.iter() if elem.tag.rsplit("}", 1)[-1] == "p"]:
        text_nodes = [elem for elem in paragraph.iter() if elem.tag.rsplit("}", 1)[-1] == "t"]
        if not text_nodes:
            continue
        joined = "".join(elem.text or "" for elem in text_nodes)
        replaced, occurrences = _replace_mapped_text(joined, replacements)
        if occurrences:
            text_nodes[0].text = replaced
            for node in text_nodes[1:]:
                node.text = ""
            count += occurrences

    for elem in root.iter():
        if elem.text:
            elem.text, occurrences = _replace_mapped_text(elem.text, replacements)
            count += occurrences
        if elem.tail:
            elem.tail, occurrences = _replace_mapped_text(elem.tail, replacements)
            count += occurrences
        for key, value in list(elem.attrib.items()):
            replaced, occurrences = _replace_mapped_text(value, replacements)
            if occurrences:
                elem.attrib[key] = replaced
                count += occurrences
    return ElementTree.tostring(root, encoding="utf-8", xml_declaration=True), count


def _extract_xml_payload_text(data: bytes) -> str:
    """Extract text and attribute values from one OOXML XML relationship part."""
    try:
        root = ElementTree.fromstring(data)
    except ElementTree.ParseError as error:
        raise ValueError("OOXML contains an invalid textual package part") from error

    values: List[str] = []

    def append_if_textual(value: str) -> None:
        candidate = value.strip()
        if candidate:
            has_letters = any(character.isalpha() for character in candidate)
            has_digits = any(character.isdigit() for character in candidate)
            has_text = (
                "@" in candidate
                or (has_letters and not has_digits)
                or (has_letters and any(character.isspace() for character in candidate))
            )
            formatted_phone = bool(PHONE_PATTERN.fullmatch(candidate)) and any(
                separator in candidate for separator in ("+", "-", "/", " ")
            )
            if has_text or formatted_phone:
                values.append(candidate)

    for elem in root.iter():
        if elem.text:
            append_if_textual(elem.text)
        if elem.tail:
            append_if_textual(elem.tail)
        for value in elem.attrib.values():
            append_if_textual(value)
    # A visible record separator prevents regexes from joining unrelated XML
    # values across attribute/element boundaries via ``\s``.
    return "\u241e".join(values)


def _xml_payload_value_text(data: bytes) -> str:
    """Return XML character data and attribute values without element names."""
    try:
        root = ElementTree.fromstring(data)
    except ElementTree.ParseError:
        return data.decode("utf-8", errors="replace")
    values: List[str] = []
    for elem in root.iter():
        if elem.text:
            values.append(elem.text)
        if elem.tail:
            values.append(elem.tail)
        values.extend(elem.attrib.values())
    return "\u241e".join(values)


def _extract_ooxml_package_text(path: Path) -> str:
    """Read every textual OOXML surface used by the later package sanitizer."""
    infos = _validate_archive(path)
    values: List[str] = []
    with zipfile.ZipFile(path, "r") as archive:
        for info in infos:
            if info.filename.casefold().endswith((".xml", ".rels")):
                extracted = _extract_xml_payload_text(archive.read(info))
                if extracted:
                    values.append(extracted)
    return "\u241e".join(values)


def _sanitize_ooxml_package(path: Path, replacements: List[Tuple[str, str]]) -> int:
    _validate_archive(path)
    count = 0
    fd, temp_name = tempfile.mkstemp(prefix="anonymizer_ooxml_", suffix=path.suffix)
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
            temp_path, "w", compression=zipfile.ZIP_DEFLATED, allowZip64=True
        ) as target:
            for info in source.infolist():
                data = source.read(info)
                lower_name = info.filename.casefold()
                if lower_name.endswith((".xml", ".rels")):
                    data, occurrences = _rewrite_xml_payload(data, replacements)
                    count += occurrences
                target.writestr(info, data)
        shutil.copy2(temp_path, path)
    finally:
        temp_path.unlink(missing_ok=True)
    return count


def _verify_ooxml_no_residuals(path: Path, replacements: List[Tuple[str, str]]) -> None:
    with zipfile.ZipFile(path, "r") as archive:
        for info in archive.infolist():
            if info.filename.casefold().endswith((".xml", ".rels")):
                text = _xml_payload_value_text(archive.read(info))
                if _residual_originals(text, replacements):
                    raise ValueError("OOXML residual verification failed")


def _anonymize_relative_path(relative_path: Path, profile: AnonymProfile) -> Path:
    replacements = _collect_replacements(profile)
    safe_parts: List[str] = []
    for part in relative_path.parts:
        replaced, _ = _replace_mapped_text(part, replacements)
        if replaced in ("", ".", "..") or Path(replaced).name != replaced:
            raise ValueError("Anonymized path component is unsafe")
        safe_parts.append(replaced)
    return Path(*safe_parts)


def _publish_tree_atomically(staging: Path, destination: Path) -> None:
    """Copy into a same-volume sibling and expose the complete tree in one rename."""
    destination.parent.mkdir(parents=True, exist_ok=True)
    _reject_symlink(destination.parent)
    _reject_symlink(destination, allow_missing_leaf=True)

    publish_root = Path(
        tempfile.mkdtemp(prefix=".anonymizer-publish-", dir=destination.parent)
    )
    candidate = publish_root / "tree"
    try:
        shutil.copytree(staging, candidate)
        if destination.exists():
            if not destination.is_dir() or any(destination.iterdir()):
                raise ValueError("Destination must remain absent or empty until publication")
            destination.rmdir()
        os.replace(candidate, destination)
    finally:
        shutil.rmtree(publish_root, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════
# Personennamen-Erkennung (spaCy NER, mehrsprachig)
# ═══════════════════════════════════════════════════════════════

# DE- und EN-Modell kombiniert: Klienten/Familien in diesem Kontext haben oft
# nicht-deutsche Namen (frankophon, afrikanisch, etc.), die ein rein deutsches
# NER-Modell seltener zuverlaessig als PERSON erkennt.
NER_MODELS = ("de_core_news_lg", "en_core_web_lg")
_NER_PERSON_LABELS = {"PER", "PERSON"}
_NER_KEEP_COMPONENTS = {"tok2vec", "ner"}

# Sicherheits-Obergrenze: Ein einzelnes klientenbezogenes Dokument (Protokoll,
# Hilfeplan, Gruppenprotokoll) erwaehnt realistisch nur eine Handvoll Personen.
# Weit mehr Treffer deuten auf generisches Referenzmaterial (Fachbuch, Spiele-
# sammlung mit zitierten Autoren) oder fehlerhafte Text-Extraktion hin -- in
# dem Fall NICHT blind Dutzende/Hunderte Fake-Namen erzeugen (Korruptions-
# risiko bei ueberlappenden Ersetzungen), sondern das Dokument ueberspringen.
_NER_MAX_NAMES_PER_CHUNK = 60

# Deutsche Fuellwoerter/Praepositionen -- tauchen sie MITTEN in einem erkannten
# "Namen" auf, ist die Entitaetsgrenze falsch (typisches Symptom bei fehl-
# uebertragenen NER-Modellen, z.B. englisches Modell auf deutschem Fliesstext).
_NER_MID_SPAN_STOPWORDS = {
    "der", "die", "das", "des", "dem", "den", "und", "oder", "im", "am",
    "zu", "zur", "zum", "auf", "in", "an", "bei", "mit", "von", "vom",
    "fuer", "für", "ueber", "über", "unter", "durch", "ohne", "gegen", "um",
    "ist", "sind", "war", "hat", "eine", "einen", "einer",
    # Reflexivpronomen/Modal- und Hilfsverben -- in echten Namen nie
    # enthalten, aber in klinischen/foerderpaedagogischen Beobachtungssaetzen
    # haeufig ("... bewegt sich grob", "... kann sich noch nicht ..."). Ohne
    # diese Erweiterung reisst ein faelschlich zu weit gefasster NER-Span
    # ganze Beobachtungsfragmente statt nur des Namens heraus.
    "sich", "ihm", "ihr", "ihn", "ihre", "ihrem", "ihren", "ihrer",
    "sein", "seine", "seinem", "seinen", "seiner",
    "wird", "wurde", "kann", "muss", "soll", "will", "werden", "worden",
    "hatte", "haette", "hätte", "waere", "wäre", "wuerde", "würde",
    "dass", "weil", "wenn", "als", "wie", "noch", "nur", "sehr",
    "man", "es", "er", "sie",
}

# Generische Rollenbegriffe (Klient/Patient/Angehoerige) -- werden von NER
# gelegentlich als PERSON erkannt, sind aber KEINE identifizierenden Namen,
# sondern gewoehnliche deutsche Substantive. Client-unabhaengig, gilt fuer
# jeden Foerderbericht.
_NER_GENERIC_ROLE_NOUNS = {
    "klient", "klienten", "klientin", "klientinnen",
    "patient", "patienten", "patientin", "patientinnen",
    "kind", "kinder", "junge", "jungen", "mädchen",
    "schüler", "schülerin", "schülerinnen",
    "mutter", "vater", "eltern", "bruder", "schwester", "geschwister",
    "therapeut", "therapeutin", "betreuer", "betreuerin",
    "lehrer", "lehrerin", "mitarbeiter", "mitarbeiterin",
}

# Verwaltungs-/Einrichtungs- und Berichts-Gattungsbegriffe -- reales
# Fehlalarm-Muster: NER verschmilzt ein Gattungswort MIT einem echten
# Eigennamen zu einem einzigen PERSON-Span (z.B. "Landkreis Loerrach",
# "Jugendamt Musterstadt") oder markiert ein einzelnes Berichts-Substantiv
# faelschlich als Namen ("Foerderung", "Zusage", "Ablauf"). Client- und
# themenunabhaengig, gilt fuer jeden Foerder-/Hilfeplanbericht.
_NER_GENERIC_REPORT_NOUNS = {
    "landkreis", "kreis", "stadt", "gemeinde", "bezirk", "region",
    "amt", "jugendamt", "landesamt", "sozialamt", "gesundheitsamt",
    "behoerde", "behörde", "dienststelle", "verwaltung", "traeger", "träger",
    "verein", "institut", "einrichtung", "gruppe", "abteilung",
    "schule", "kita", "kindergarten", "klasse",
    "foerderung", "förderung", "zusage", "ablauf", "bericht", "termin",
    "gespraech", "gespräch", "sitzung", "protokoll",
    "hilfeplan", "foerderplan", "förderplan", "ziel", "ziele",
    "massnahme", "maßnahme", "massnahmen", "maßnahmen",
    "entwicklung", "verlauf", "stand", "beginn", "ende", "zeitraum",
    "datum", "unterschrift", "anlage", "anhang",
}


def _looks_like_person_name(name: str) -> bool:
    """
    Grobe Plausibilitaetsprüfung fuer einen von NER erkannten "Personennamen".

    Faengt falsche Entitaetsgrenzen ab (z.B. wenn ein Modell eine ganze
    mehrzeilige Phrase oder einen zusammengesetzten Fachbegriff faelschlich
    als EIN Name erkennt) -- typisches Symptom, wenn das englische Modell
    auf deutschem Fliesstext angewendet wird.
    """
    name = name.strip()
    if not name or len(name) > 40 or "\n" in name or "\t" in name:
        return False
    words = name.split()
    if not (1 <= len(words) <= 4):
        return False
    cleaned_words = []
    for word in words:
        cleaned = word.strip(".,;:!?()[]{}\"'-")
        if not cleaned:
            return False
        if cleaned.lower() in _NER_MID_SPAN_STOPWORDS:
            return False
        if not cleaned[0].isupper():
            return False
        cleaned_words.append(cleaned)
    # Gattungsbegriffe an JEDER Wortposition ablehnen (nicht nur bei
    # Einzelwort-Treffern) -- NER verschmilzt ein Gattungswort haeufig mit
    # einem echten Eigennamen zu einem Span ("Landkreis Loerrach").
    if any(
        word.lower() in _NER_GENERIC_ROLE_NOUNS
        or word.lower() in _NER_GENERIC_REPORT_NOUNS
        for word in cleaned_words
    ):
        return False
    return True


_spacy_model_cache: Dict[str, "object"] = {}


def _get_spacy_model(model_name: str):
    """Laedt ein spaCy-Modell einmalig (teuer, ~5s) und cached es pro Prozess.

    Unnoetige Pipeline-Komponenten (Parser/Tagger/Lemmatizer/...) werden beim
    Laden ausgeschlossen -- nur tok2vec (Wortvektoren) + ner werden fuer die
    Personennamen-Erkennung gebraucht.
    """
    if model_name in _spacy_model_cache:
        return _spacy_model_cache[model_name]
    if not SPACY_AVAILABLE:
        return None
    try:
        full_meta = spacy.load(model_name, exclude=[])
        exclude = [name for name in full_meta.pipe_names if name not in _NER_KEEP_COMPONENTS]
        nlp = spacy.load(model_name, exclude=exclude) if exclude else full_meta
    except Exception:
        # Der aufrufende Standard-Scan entscheidet fail-closed. Die öffentliche
        # Warnung enthält bewusst keine rohen Loader-/Pfaddetails.
        print(f"[WARN] spaCy-Modell '{model_name}' konnte nicht geladen werden")
        nlp = None
    _spacy_model_cache[model_name] = nlp
    return nlp


def detect_person_names_ner(
    text: str,
    whitelist: Optional[List[str]] = None,
    *,
    fail_on_ambiguous: bool = False,
) -> List[str]:
    """
    Erkennt Personennamen im Text via spaCy-NER (DE+EN kombiniert).

    Im Unterschied zu Regex-Heuristiken kann NER PERSON von Fachbegriffen,
    Orten und Institutionen unterscheiden -- wichtig im Deutschen, wo JEDES
    Substantiv grossgeschrieben wird. Erkennt auch unbekannte Namen und
    Schreibvarianten, die keine manuell gepflegte Liste je abdecken koennte.

    Args:
        text: Zu scannender Text
        whitelist: Namen, die trotz Erkennung NICHT zurueckgegeben werden
                   (z.B. Therapeut/Amtspersonen)

    Returns:
        Sortierte Liste eindeutiger erkannter Personennamen
    """
    if not SPACY_AVAILABLE or not text.strip():
        return []

    wl_lower = {w.lower() for w in (whitelist or [])}
    found = set()

    for model_name in NER_MODELS:
        nlp = _get_spacy_model(model_name)
        if nlp is None:
            continue
        # spaCy-Modelle haben ein Zeichenlimit (Default 1_000_000) -- bei sehr
        # langen Bundles in Bloecken verarbeiten, um Fehler zu vermeiden.
        max_len = nlp.max_length
        for offset in range(0, len(text), max_len):
            chunk = text[offset:offset + max_len]
            doc = nlp(chunk)
            chunk_names = set()
            for ent in doc.ents:
                if ent.label_ not in _NER_PERSON_LABELS:
                    continue
                name = ent.text.strip()
                if not name or any(ch.isdigit() for ch in name):
                    continue
                if name.lower() in wl_lower:
                    continue
                if not _looks_like_person_name(name):
                    continue
                # Wortgrenzen-Check: folgt direkt (ohne Trennzeichen) ein
                # Kleinbuchstabe, ist die Entitaet nur der ANFANG eines
                # zusammengesetzten deutschen Wortes (z.B. NER erkennt
                # "Wahrnehmung" als Namensbeginn von "Wahrnehmungsbesonder-
                # heiten") -- Ersetzung wuerde ein Wortfragment hinterlassen.
                next_char = chunk[ent.end_char:ent.end_char + 1]
                if next_char and next_char.isalpha() and next_char.islower():
                    continue
                chunk_names.add(name)

            if len(chunk_names) > _NER_MAX_NAMES_PER_CHUNK:
                if fail_on_ambiguous:
                    raise RuntimeError(
                        "NER result exceeds the reviewed per-section identity limit"
                    )
                print(
                    f"[WARN] NER ({model_name}): {len(chunk_names)} Personennamen in einem "
                    f"Textabschnitt gefunden (> {_NER_MAX_NAMES_PER_CHUNK}) -- vermutlich "
                    f"generisches Referenzmaterial statt Klientendaten, Abschnitt wird "
                    f"NICHT automatisch anonymisiert (Korruptionsrisiko)."
                )
                continue
            found.update(chunk_names)

    return sorted(found)


# ═══════════════════════════════════════════════════════════════
# Lokaler Schlüssel-Speicher (NICHT in OneDrive)
# ═══════════════════════════════════════════════════════════════

def get_local_keys_dir() -> Path:
    """
    Gibt den lokalen Schlüssel-Ordner zurück (außerhalb OneDrive).

    Schlüssel-Dateien (.schluessel.enc) dürfen NICHT über Cloud-Dienste
    synchronisiert werden.

    Auflösung:
        1. ANONYMIZER_KEYS_DIR (ENV) — nach lokaler Pfadprüfung nutzen
        2. Windows: %LOCALAPPDATA%\\anonymizer\\keys
        3. Unix   : ~/.local/share/anonymizer/keys

    Returns:
        Path zum lokalen keys-Ordner
    """
    env_override = os.environ.get("ANONYMIZER_KEYS_DIR")
    if env_override:
        keys_dir = Path(env_override).expanduser()
        if _is_cloud_path(keys_dir):
            raise ValueError("ANONYMIZER_KEYS_DIR must not point to a cloud-synced location")
        _reject_symlink(keys_dir, allow_missing_leaf=True)
        keys_dir.mkdir(parents=True, exist_ok=True, mode=0o700)
        _reject_symlink(keys_dir)
        try:
            keys_dir.chmod(0o700)
        except OSError:
            pass
        return keys_dir.resolve(strict=True)

    if os.name == "nt":
        base = Path(os.environ.get("LOCALAPPDATA", os.path.expanduser("~")))
    else:
        base = Path.home() / ".local" / "share"
    keys_dir = base / "anonymizer" / "keys"
    if _is_cloud_path(keys_dir):
        raise ValueError("Resolved key directory is cloud-synced")
    _reject_symlink(keys_dir, allow_missing_leaf=True)
    keys_dir.mkdir(parents=True, exist_ok=True, mode=0o700)
    _reject_symlink(keys_dir)
    try:
        keys_dir.chmod(0o700)
    except OSError:
        pass
    return keys_dir.resolve(strict=True)


def get_key_path(client_id: str) -> Path:
    """
    Gibt den sicheren Pfad für eine Schlüssel-Datei zurück.

    Args:
        client_id: Klienten-ID (z.B. 'K_0042')

    Returns:
        Path zur .schluessel.enc Datei im lokalen Speicher
    """
    safe_client_id = _validate_client_id(client_id)
    keys_dir = get_local_keys_dir()
    result = keys_dir / f"{safe_client_id}.schluessel.enc"
    if result.parent.resolve(strict=True) != keys_dir.resolve(strict=True):
        raise ValueError("Key path escaped the local key directory")
    return result


def _clear_hidden_attribute(filepath: Path):
    """Entfernt das Windows Hidden-Attribut (nötig vor Überschreiben)."""
    if os.name == "nt" and filepath.exists():
        try:
            import ctypes
            attrs = ctypes.windll.kernel32.GetFileAttributesW(str(filepath))
            if attrs != -1 and (attrs & 0x02):
                ctypes.windll.kernel32.SetFileAttributesW(str(filepath), attrs & ~0x02)
        except Exception:
            pass


def _set_hidden_attribute(filepath: Path):
    """Setzt das Windows Hidden-Attribut auf eine Datei (nur Windows)."""
    if os.name == "nt":
        try:
            import ctypes
            FILE_ATTRIBUTE_HIDDEN = 0x02
            ctypes.windll.kernel32.SetFileAttributesW(str(filepath), FILE_ATTRIBUTE_HIDDEN)
        except Exception:
            pass  # Nicht kritisch


# ═══════════════════════════════════════════════════════════════
# AES-Verschlüsselung für Schlüssel
# ═══════════════════════════════════════════════════════════════

def _derive_key(password: str, salt: bytes) -> bytes:
    """Leitet einen Fernet-Schlüssel aus einem Passwort ab (PBKDF2)."""
    if not CRYPTO_AVAILABLE:
        raise RuntimeError("cryptography-Paket nicht installiert: pip install cryptography")

    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=480000,
    )
    return base64.urlsafe_b64encode(kdf.derive(password.encode("utf-8")))


def encrypt_key_file(profile: AnonymProfile, output_path: str, password: str) -> Path:
    """
    Speichert das AnonymProfile authentifiziert mit Fernet verschlüsselt.

    Args:
        profile: Das Anonymisierungsprofil mit allen Mappings
        output_path: Pfad für die verschlüsselte Datei (.schluessel.enc)
        password: Passwort für die Verschlüsselung
    """
    if not CRYPTO_AVAILABLE:
        raise RuntimeError("cryptography-Paket nicht installiert: pip install cryptography")
    if not isinstance(password, str) or len(password) < 12:
        raise ValueError("Password must contain at least 12 characters")

    # Profil als JSON serialisieren
    data = {
        "version": profile.version,
        "client_id": profile.client_id,
        "tarnname": profile.tarnname,
        "fake_geburtsdatum": profile.fake_geburtsdatum,
        "mappings": profile.mappings,
        "created": profile.created or datetime.now().isoformat()
    }
    plaintext = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")

    # Verschlüsseln
    salt = os.urandom(16)
    key = _derive_key(password, salt)
    fernet = Fernet(key)
    encrypted = fernet.encrypt(plaintext)

    # Speichern: Salt (16 Bytes fest) + verschlüsselte Daten
    path = Path(output_path).expanduser()
    if _is_cloud_path(path):
        raise ValueError("Encrypted key files must not be written to cloud-synced paths")
    _reject_symlink(path, allow_missing_leaf=True)
    path.parent.mkdir(parents=True, exist_ok=True, mode=0o700)
    _reject_symlink(path.parent)
    _clear_hidden_attribute(path)  # Falls Datei existiert und Hidden ist

    temp_path = path.parent / f".{path.name}.{secrets.token_hex(8)}.tmp"
    flags = os.O_WRONLY | os.O_CREAT | os.O_EXCL
    if hasattr(os, "O_BINARY"):
        flags |= os.O_BINARY
    fd = os.open(temp_path, flags, 0o600)
    try:
        with os.fdopen(fd, "wb") as stream:
            stream.write(salt)  # genau 16 Bytes
            stream.write(encrypted)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temp_path, path)
        try:
            path.chmod(0o600)
        except OSError:
            pass
    finally:
        temp_path.unlink(missing_ok=True)

    # Hidden-Attribut setzen (Schutz vor versehentlichem Zugriff)
    _set_hidden_attribute(path)

    return path


def decrypt_key_file(key_path: str, password: str) -> AnonymProfile:
    """
    Lädt und entschlüsselt ein AnonymProfile.

    Args:
        key_path: Pfad zur .schluessel.enc Datei
        password: Passwort für die Entschlüsselung
    """
    if not CRYPTO_AVAILABLE:
        raise RuntimeError("cryptography-Paket nicht installiert: pip install cryptography")

    path = Path(key_path).expanduser()
    if _is_cloud_path(path):
        raise ValueError("Encrypted key files must not be read from cloud-synced paths")
    path = _validate_regular_input(path)
    if path.stat().st_size > _MAX_KEY_FILE_BYTES:
        raise ValueError("Key file exceeds the 10 MiB safety limit")
    with open(path, "rb") as f:
        content = f.read(_MAX_KEY_FILE_BYTES + 1)
    if not 32 <= len(content) <= _MAX_KEY_FILE_BYTES:
        raise ValueError("Invalid encrypted key file length")

    # Salt (erste 16 Bytes) und verschlüsselte Daten trennen
    salt = content[:16]
    encrypted = content[16:]

    # Entschlüsseln
    key = _derive_key(password, salt)
    fernet = Fernet(key)
    plaintext = fernet.decrypt(encrypted)

    # JSON parsen
    data = json.loads(plaintext.decode("utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Invalid encrypted profile payload")
    for required in ("client_id", "tarnname", "fake_geburtsdatum", "mappings"):
        if required not in data:
            raise ValueError("Invalid encrypted profile payload")
    _validate_client_id(data["client_id"])
    if not isinstance(data["mappings"], dict):
        raise ValueError("Invalid encrypted profile mappings")

    return AnonymProfile(
        client_id=data["client_id"],
        tarnname=data["tarnname"],
        fake_geburtsdatum=data["fake_geburtsdatum"],
        mappings=data.get("mappings", {}),
        created=data.get("created", ""),
        version=data.get("version", 1)
    )


def _extract_legacy_doc_text(filepath: str) -> str:
    """
    Extrahiert Text aus altem Word-Binaerformat (.doc) via antiword oder
    LibreOffice (soffice). Spiegelt dieselbe Fallback-Kette wie
    document_pipeline.py::_extract_doc(), damit Sensible-Daten-Scan und
    LLM-Prompt-Buendelung denselben Dokumenteninhalt sehen.
    """
    import subprocess

    source = _validate_regular_input(Path(filepath))
    if source.stat().st_size > 25 * 1024 * 1024:
        raise ValueError("Legacy DOC exceeds the 25 MiB conversion limit")

    if shutil.which("antiword"):
        try:
            result = subprocess.run(
                ["antiword", str(source)],
                capture_output=True, text=True,
                encoding="utf-8", errors="replace", timeout=30,
                stdin=subprocess.DEVNULL,
            )
            if result.returncode == 0 and result.stdout.strip():
                if len(result.stdout.encode("utf-8")) > 50 * 1024 * 1024:
                    raise ValueError("Legacy DOC text exceeds the conversion limit")
                return result.stdout.strip()
        except Exception:
            pass

    if shutil.which("soffice"):
        try:
            import tempfile
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run(
                    ["soffice", "--headless", "--convert-to", "txt:Text",
                     "--outdir", tmpdir, str(source)],
                    capture_output=True, timeout=60, stdin=subprocess.DEVNULL,
                )
                if result.returncode == 0:
                    txt_file = Path(tmpdir) / (source.stem + ".txt")
                    if txt_file.exists():
                        if txt_file.stat().st_size > 50 * 1024 * 1024:
                            raise ValueError("Legacy DOC text exceeds the conversion limit")
                        return txt_file.read_text(encoding="utf-8", errors="replace")
        except Exception:
            pass

    return ""


# ═══════════════════════════════════════════════════════════════
# Anonymisierer
# ═══════════════════════════════════════════════════════════════

class DocumentAnonymizer:
    """
    Anonymisiert Dokumente durch konsistente Pseudonymisierung.

    Workflow:
        1. create_profile() - Profil mit Tarnnamen erstellen
        2. anonymize_folder() - Alle Dokumente im Ordner anonymisieren
        3. Schlüssel wird authentifiziert verschlüsselt gespeichert
    """

    def __init__(
        self,
        *,
        require_ner: bool = True,
        allow_unverified_media: bool = False,
        trusted_template_path: Optional[str] = None,
    ):
        self._progress = ProgressInfo()
        self._used_names: set = set()
        self.require_ner = require_ner
        self.allow_unverified_media = allow_unverified_media
        # ENV-Fallback nur wenn kein expliziter Parameter uebergeben wurde.
        resolved_template = trusted_template_path or os.environ.get(
            "ANONYMIZER_TRUSTED_TEMPLATE"
        )
        self.trusted_media_hashes = _load_trusted_template_media_hashes(resolved_template)
        self.global_whitelist = self._load_global_whitelist()

    def _load_global_whitelist(self) -> dict:
        """
        Lädt die optionale globale Whitelist.

        Sucht unter ANONYMIZER_HOME (ENV) → ~/.anonymizer/anonymizer_whitelist.json.
        """
        anon_home = Path(os.environ.get("ANONYMIZER_HOME", Path.home() / ".anonymizer"))
        whitelist_file = anon_home / "anonymizer_whitelist.json"
        if whitelist_file.exists():
            try:
                return json.loads(whitelist_file.read_text(encoding="utf-8"))
            except Exception:
                print("[WARN] Anonymizer Whitelist konnte nicht geladen werden")
        return {"titles": [], "names": [], "organizations": []}

    @property
    def progress(self) -> ProgressInfo:
        return self._progress

    def scan_text_for_sensitive_data(self, text: str) -> Dict[str, List[str]]:
        """
        Scannt Text nach sensiblen Daten (Telefon, E-Mail, Adressen, Institutionen).

        Args:
            text: Der zu scannende Text

        Returns:
            Dict mit Listen: {"phones": [...], "emails": [...], "addresses": [...], "institutions": [...]}
        """
        found = {
            "phones": [],
            "emails": [],
            "addresses": [],
            "institutions": [],
            "table_row_names": [],
            "ner_person_names": []
        }

        # Telefonnummern finden
        phones = PHONE_PATTERN.findall(text)
        for phone in phones:
            cleaned = phone.strip()
            if cleaned and cleaned not in found["phones"]:
                found["phones"].append(cleaned)

        # Jede E-Mail-Adresse ist personenbezogen, unabhaengig von der Domain.
        emails = EMAIL_PATTERN.findall(text)
        seen_email_keys = set()
        for email in emails:
            cleaned = email.strip()
            dedupe_key = cleaned.casefold()
            if cleaned and dedupe_key not in seen_email_keys:
                found["emails"].append(cleaned)
                seen_email_keys.add(dedupe_key)

        # Adressen finden
        addresses = STREET_PATTERN.findall(text)
        for addr in addresses:
            cleaned = addr.strip()
            if cleaned and cleaned not in found["addresses"]:
                found["addresses"].append(cleaned)

        # Institutionsnamen finden (Schulen, Kliniken, etc.)
        institutions = INSTITUTION_PATTERN.findall(text)
        for inst in institutions:
            cleaned = inst.strip()
            if cleaned and cleaned not in found["institutions"]:
                # Nicht auf globaler Whitelist?
                if not any(w.lower() in cleaned.lower() for w in self.global_whitelist.get("organizations", [])):
                    found["institutions"].append(cleaned)

        # Personennamen in Tabellenzeilen (z.B. Teilnehmerlisten in Gruppenprotokollen)
        table_names = TABLE_ROW_NAME_PATTERN.findall(text)
        for name in table_names:
            cleaned = name.strip()
            if cleaned and cleaned not in found["table_row_names"]:
                if not any(w.lower() == cleaned.lower() for w in self.global_whitelist.get("names", [])):
                    found["table_row_names"].append(cleaned)

        # Generische Personennamen-Erkennung (spaCy NER, DE+EN) -- Hauptmechanismus
        # fuer unbekannte Drittpersonen; faengt auch Namen und Schreibvarianten,
        # die keine Regel/Liste explizit vorgesehen hat.
        if SPACY_AVAILABLE:
            model_available = any(_get_spacy_model(name) is not None for name in NER_MODELS)
            if self.require_ner and not model_available:
                raise RuntimeError(
                    "Person-name scanning requires at least one configured spaCy NER model; "
                    "set require_ner=False only for an explicitly reviewed reduced-coverage workflow"
                )
            ner_names = detect_person_names_ner(
                text,
                whitelist=self.global_whitelist.get("names", []),
                fail_on_ambiguous=self.require_ner,
            )
            found["ner_person_names"] = ner_names
        elif self.require_ner:
            raise RuntimeError(
                "Person-name scanning requires spaCy and a configured NER model; "
                "set require_ner=False only for an explicitly reviewed reduced-coverage workflow"
            )

        return found

    def create_profile(
        self,
        real_name: str,
        geburtsdatum: str,
        weitere_namen: Optional[List[str]] = None,
        weitere_daten: Optional[Dict[str, str]] = None,
        whitelist: Optional[List[str]] = None,
        scanned_data: Optional[Dict[str, List[str]]] = None
    ) -> AnonymProfile:
        """
        Erstellt ein Anonymisierungsprofil.

        Args:
            real_name: Echter Name des Klienten (z.B. "Max Mustermann")
            geburtsdatum: Echtes Geburtsdatum (dd.mm.yyyy)
            weitere_namen: Weitere zu ersetzende Namen (Ärzte, Angehörige)
            weitere_daten: Weitere Daten {"adresse": "Musterstr. 5", "telefon": "07761/123"}
            whitelist: Namen die NICHT anonymisiert werden (z.B. Sachbearbeiter vom Amt)
            scanned_data: Automatisch erkannte Daten aus scan_text_for_sensitive_data()
                          {"phones": [...], "emails": [...], "addresses": [...]}
        """
        # Whitelist zusammenstellen (Global + Parameter)
        whitelist_set = set(whitelist) if whitelist else set()
        whitelist_set.update(self.global_whitelist.get("names", []))
        whitelist_set.update(self.global_whitelist.get("organizations", []))

        # Titel für automatische Erkennung
        titles = self.global_whitelist.get("titles", [])
        # Client-ID generieren
        client_id = f"K_{secrets.token_hex(3).upper()}"

        # Vorname extrahieren für Gender-Erkennung
        real_parts = real_name.strip().split()
        original_vorname = real_parts[0] if real_parts else ""

        # Tarnname mit Geschlechts-Matching generieren
        tarnname = _generate_tarnname(
            used_names=self._used_names,
            original_vorname=original_vorname
        )
        self._used_names.add(tarnname)

        # Falsches Geburtsdatum (gleiches Alter)
        fake_geb, date_offset = _generate_fake_date_same_age(geburtsdatum)

        # Mappings aufbauen
        mappings: Dict[str, Dict[str, str]] = {
            "names": {},
            "dates": {},
            "addresses": {},
            "phones": {},
            "emails": {},
            "misc": {}
        }

        # Hauptname
        real_parts = real_name.strip().split()
        tarn_parts = tarnname.strip().split()

        mappings["names"][real_name] = tarnname
        # Auch Vorname und Nachname einzeln
        if len(real_parts) >= 2 and len(tarn_parts) >= 2:
            mappings["names"][real_parts[0]] = tarn_parts[0]   # Vorname
            mappings["names"][real_parts[-1]] = tarn_parts[-1]  # Nachname

        # Weitere Namen (nur wenn nicht auf Whitelist und keine Amtsperson)
        if weitere_namen:
            for name in weitere_namen:
                # Check ob explizit whitelisted
                if name in whitelist_set:
                    continue

                # Check ob Amtsperson durch Titel (z.B. "Dr. Meyer" -> "Dr." ist Title)
                is_amtsperson = any(title in name for title in titles)
                if is_amtsperson:
                    continue

                fake = _generate_tarnname(self._used_names)
                self._used_names.add(fake)
                mappings["names"][name] = fake

        # Geburtsdatum
        mappings["dates"][geburtsdatum] = fake_geb

        # Weitere Daten (explizit übergeben)
        if weitere_daten:
            if "adresse" in weitere_daten:
                mappings["addresses"][weitere_daten["adresse"]] = _generate_fake_address()
            if "telefon" in weitere_daten:
                mappings["phones"][weitere_daten["telefon"]] = _generate_fake_phone()
            if "email" in weitere_daten:
                mappings["emails"][weitere_daten["email"]] = _generate_fake_email()
            # Restliche als misc
            for key, val in weitere_daten.items():
                if key not in ("adresse", "telefon", "email") and val:
                    mappings["misc"][val] = f"[{key.upper()}_ANON]"

        # Automatisch gescannte Daten (Telefon, E-Mail, Adressen)
        if scanned_data:
            # Telefonnummern
            for phone in scanned_data.get("phones", []):
                if phone not in mappings["phones"]:
                    mappings["phones"][phone] = _generate_fake_phone()

            # E-Mail-Adressen
            for email in scanned_data.get("emails", []):
                if email not in mappings["emails"]:
                    mappings["emails"][email] = _generate_fake_email(email)

            # Straßenadressen
            for addr in scanned_data.get("addresses", []):
                if addr not in mappings["addresses"]:
                    mappings["addresses"][addr] = _generate_fake_address()

            # Institutionsnamen (Schulen, Kliniken, etc.)
            for inst in scanned_data.get("institutions", []):
                if inst not in mappings.get("institutions", {}):
                    if "institutions" not in mappings:
                        mappings["institutions"] = {}
                    mappings["institutions"][inst] = _generate_fake_institution(inst)

            # Alle automatisch erkannten Personennamen muessen auch in das
            # Profil eingehen; Erkennung ohne Mapping waere ein Fail-open.
            for detected_name in (
                scanned_data.get("table_row_names", [])
                + scanned_data.get("ner_person_names", [])
            ):
                if not detected_name or detected_name.casefold() in {
                    item.casefold() for item in whitelist_set
                }:
                    continue
                if detected_name in mappings["names"]:
                    continue
                fake = _generate_tarnname(self._used_names)
                self._used_names.add(fake)
                mappings["names"][detected_name] = fake

        return AnonymProfile(
            client_id=client_id,
            tarnname=tarnname,
            fake_geburtsdatum=fake_geb,
            mappings=mappings,
            created=datetime.now().isoformat()
        )

    def extract_text_from_file(self, filepath: str) -> str:
        """
        Extrahiert Text aus einer Datei zum Scannen.

        Args:
            filepath: Pfad zur Datei

        Returns:
            Extrahierter Text
        """
        path = Path(filepath)
        suffix = path.suffix.lower()
        text = ""

        try:
            _validate_regular_input(path)
            if suffix == ".docx":
                if not DOCX_AVAILABLE:
                    raise RuntimeError("python-docx is required for DOCX scanning")
                doc = Document(str(path))
                paragraphs = [p.text for p in doc.paragraphs]
                # Auch Tabellen
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs.append(cell.text)
                # The package-level pass adds comments, headers/footers,
                # footnotes, text boxes, properties, hyperlinks, and custom XML.
                paragraphs.append("\u241e" + _extract_ooxml_package_text(path) + "\u241e")
                text = "\n".join(paragraphs)

            elif suffix == ".doc":
                # Altes Word-Binaerformat -- python-docx kann das NICHT lesen.
                # Ohne diesen Zweig bleiben .doc-Dateien (z.B. Aktendeckblatt)
                # bei der Sensible-Daten-Suche komplett unerkannt (leerer Text),
                # obwohl ihr Inhalt sehr wohl in den LLM-Prompt gebuendelt wird
                # (document_pipeline.py nutzt dieselbe Extraktion fuers Bundling).
                text = _extract_legacy_doc_text(str(path))

            elif suffix in (".txt", ".md"):
                try:
                    text = path.read_text(encoding="utf-8")
                except UnicodeDecodeError:
                    text = path.read_text(encoding="latin-1")

            elif suffix == ".pdf":
                if not FITZ_AVAILABLE:
                    raise RuntimeError("PyMuPDF is required for PDF scanning")
                doc = fitz.open(str(path))
                for page in doc:
                    text += page.get_text()
                doc.close()

            elif suffix == ".xlsx":
                if not EXCEL_AVAILABLE:
                    raise RuntimeError("openpyxl is required for XLSX scanning")
                wb = openpyxl.load_workbook(str(path), data_only=True, keep_links=False)
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value:
                                text += str(cell.value) + " "
                wb.close()
                # Include comments/authors, headers/footers, sheet/workbook
                # metadata, hyperlinks/relationships, and every other XML part.
                text += "\u241e" + _extract_ooxml_package_text(path) + "\u241e"
            else:
                raise ValueError("Unsupported input format")
        except Exception as e:
            raise RuntimeError(f"Text extraction failed for {suffix or 'unknown'} input") from e

        return text

    def scan_folder_for_sensitive_data(self, folder: str) -> Dict[str, List[str]]:
        """
        Scannt ALLE Dateien in einem Ordner (rekursiv) nach sensiblen Daten.

        VORSICHT bei generischen Klienten-Ordnern: Enthaelt der Ordner neben
        klientenbezogenen Dokumenten auch generische Referenz-/Methodenmaterialien
        (z.B. Spielesammlungen, Fachbuecher), werden auch DEREN Personennamen
        (z.B. zitierte Autoren) als "Drittpersonen" erkannt -- das ist meist
        nicht gewollt. Fuer Klienten-Akten mit CORE/STUFE2/EXTENDED-Struktur
        `scan_files_for_sensitive_data()` mit einer vorgefilterten Dateiliste
        bevorzugen (siehe ReportWorkflowService.create_temp_profile).

        Args:
            folder: Pfad zum Ordner

        Returns:
            Aggregierte gefundene Daten {"phones": [...], "emails": [...], "addresses": [...]}
        """
        src = Path(folder)
        _reject_symlink(src)
        if not src.is_dir():
            raise ValueError("Scan source must be an existing directory")
        source_root = src.resolve(strict=True)
        filepaths: List[Path] = []
        for root, directories, names in os.walk(source_root, followlinks=False):
            root_path = Path(root)
            for directory in list(directories):
                if _is_link_like(root_path / directory):
                    raise ValueError("Scan source contains a link-like directory")
            for name in names:
                candidate = root_path / name
                if _is_link_like(candidate):
                    raise ValueError("Scan source contains a link-like file")
                if candidate.suffix.lower() in SUPPORTED_SUFFIXES:
                    filepaths.append(candidate)
        return self.scan_files_for_sensitive_data(filepaths)

    def scan_files_for_sensitive_data(self, filepaths: List[Path]) -> Dict[str, List[str]]:
        """
        Scannt eine EXPLIZITE Liste von Dateien nach sensiblen Daten (statt
        blind einen ganzen Ordnerbaum zu durchsuchen). Damit lassen sich z.B.
        generische Referenzmaterialien gezielt von der Personennamen-Erkennung
        ausschliessen.

        Args:
            filepaths: Liste von Dateipfaden

        Returns:
            Aggregierte gefundene Daten {"phones": [...], "emails": [...], "addresses": [...], ...}
        """
        combined = {
            "phones": [],
            "emails": [],
            "addresses": [],
            "institutions": [],
            "table_row_names": [],
            "ner_person_names": []
        }

        for filepath in filepaths:
            filepath = Path(filepath)
            _reject_symlink(filepath)
            if filepath.is_file() and filepath.suffix.lower() in SUPPORTED_SUFFIXES:
                text = self.extract_text_from_file(str(filepath))
                found = self.scan_text_for_sensitive_data(text)

                for key in combined:
                    for item in found.get(key, []):
                        if item not in combined[key]:
                            combined[key].append(item)

        return combined

    def anonymize_file(self, filepath: str, profile: AnonymProfile) -> Tuple[bool, int]:
        """
        Anonymisiert eine einzelne Datei transaktional.

        Returns:
            (success, replacement_count)
        """
        source = Path(filepath)
        suffix = source.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            return False, 0
        if _residual_originals(source.name, _collect_replacements(profile)):
            # Die Einzeldatei-API kann keinen neuen Pfad zurückgeben. Ein
            # identifizierender Dateiname wird deshalb nicht als Erfolg
            # ausgegeben; dafür ist der ordnerweite Workflow vorgesehen.
            return False, 0
        try:
            source = _validate_regular_input(source)
        except (OSError, ValueError):
            return False, 0

        with tempfile.TemporaryDirectory(prefix="anonymizer_work_") as temp_dir:
            work = Path(temp_dir) / f"input{suffix}"
            shutil.copyfile(source, work)
            try:
                success, count = self._anonymize_file_in_place(work, profile)
                if not success:
                    return False, 0

                processed = work.with_suffix(".txt") if suffix == ".doc" else work
                if not processed.is_file():
                    return False, 0
                final_path = source.with_suffix(".txt") if suffix == ".doc" else source
                _reject_symlink(final_path, allow_missing_leaf=True)
                fd, sanitized_name = tempfile.mkstemp(
                    prefix=f".{final_path.name}.", suffix=".sanitized", dir=final_path.parent
                )
                os.close(fd)
                sanitized_path = Path(sanitized_name)
                try:
                    shutil.copyfile(processed, sanitized_path)
                    with open(sanitized_path, "rb+") as stream:
                        os.fsync(stream.fileno())
                    os.replace(sanitized_path, final_path)
                    if final_path != source:
                        source.unlink(missing_ok=True)
                finally:
                    sanitized_path.unlink(missing_ok=True)
                return True, count
            except Exception:
                return False, 0

    def _anonymize_file_in_place(
        self, path: Path, profile: AnonymProfile
    ) -> Tuple[bool, int]:
        suffix = path.suffix.lower()
        if suffix == ".docx":
            return self._anonymize_docx(path, profile)
        if suffix in (".txt", ".md"):
            return self._anonymize_text(path, profile)
        if suffix == ".pdf":
            return self._anonymize_pdf(path, profile)
        if suffix == ".xlsx":
            return self._anonymize_excel(path, profile)
        if suffix == ".doc":
            return self._anonymize_doc(path, profile)
        return False, 0

    def anonymize_folder(
        self,
        folder: str,
        profile: AnonymProfile,
        password: str,
        output_folder: Optional[str] = None
    ) -> AnonymResult:
        """
        Anonymisiert alle Dokumente in einem Ordner.

        Args:
            folder: Quellordner
            profile: Anonymisierungsprofil
            password: Passwort für den Schlüssel
            output_folder: Zielordner (default: ANONYMIZER_HOME/klienten/<client_id>)
        """
        result = AnonymResult()
        src = Path(folder)

        if output_folder:
            dest = Path(output_folder)
        else:
            anon_home = Path(os.environ.get("ANONYMIZER_HOME", Path.home() / ".anonymizer"))
            dest = anon_home / "klienten" / profile.client_id

        _reject_symlink(src)
        if not src.is_dir():
            raise ValueError("Source folder must be an existing directory")
        source_root = src.resolve(strict=True)
        destination = dest.expanduser().resolve(strict=False)
        if _residual_originals(str(destination), _collect_replacements(profile)):
            raise ValueError("Destination path contains a mapped identity")
        if destination == source_root or _is_relative_to(destination, source_root):
            raise ValueError("Destination must be outside the source tree")
        _reject_symlink(dest, allow_missing_leaf=True)
        if dest.exists() and any(dest.iterdir()):
            raise ValueError("Destination must be absent or empty")

        files: List[Path] = []
        unsafe_entries: List[str] = []
        for root, directories, names in os.walk(source_root, followlinks=False):
            root_path = Path(root)
            for directory in list(directories):
                if _is_link_like(root_path / directory):
                    unsafe_entries.append("link-like directory")
                    directories.remove(directory)
            for name in names:
                if name.startswith("."):
                    continue
                candidate = root_path / name
                if _is_link_like(candidate):
                    unsafe_entries.append("link-like file")
                else:
                    files.append(candidate)

        self._progress = ProgressInfo(total_files=len(files), status="anonymizing")
        with tempfile.TemporaryDirectory(prefix="anonymizer_folder_") as staging_dir:
            staging = Path(staging_dir) / "publish"
            staging.mkdir()
            seen_destinations: set[Path] = set()

            for index, filepath in enumerate(files, start=1):
                self._progress.current_file = f"file-{index}"
                self._progress.processed_files += 1
                result.processed_files += 1
                suffix = filepath.suffix.lower()
                if suffix not in SUPPORTED_SUFFIXES:
                    result.skipped_files += 1
                    result.errors.append(f"file-{index}: unsupported format")
                    continue
                try:
                    rel = filepath.relative_to(source_root)
                    if suffix == ".doc":
                        rel = rel.with_suffix(".txt")
                    safe_rel = _anonymize_relative_path(rel, profile)
                    if safe_rel in seen_destinations:
                        raise ValueError("anonymized path collision")
                    seen_destinations.add(safe_rel)

                    raw_work = Path(staging_dir) / f"raw-{index}{suffix}"
                    shutil.copyfile(filepath, raw_work)
                    success, count = self.anonymize_file(str(raw_work), profile)
                    processed = raw_work.with_suffix(".txt") if suffix == ".doc" else raw_work
                    if not success or not processed.is_file():
                        raise ValueError("format handler did not produce a verified output")
                    staged_file = staging / safe_rel
                    staged_file.parent.mkdir(parents=True, exist_ok=True)
                    shutil.move(str(processed), staged_file)
                    result.anonymized_files += 1
                    result.replacements_total += count
                except Exception:
                    result.skipped_files += 1
                    result.errors.append(f"file-{index}: anonymization failed")

            if unsafe_entries:
                result.errors.extend(unsafe_entries)
            if result.errors:
                result.anonymized_files = 0
                result.replacements_total = 0
                result.errors.append("No output was published because validation failed")
                self._progress.status = "error"
                return result

            key_path = get_key_path(profile.client_id)
            encrypt_key_file(profile, str(key_path), password)
            profil_info = {
                "client_id": profile.client_id,
                "tarnname": profile.tarnname,
                "fake_geburtsdatum": profile.fake_geburtsdatum,
                "created": profile.created,
                "files_anonymized": result.anonymized_files,
                "key_info": "Encrypted key is stored in the configured local key directory",
            }
            (staging / ".profil.json").write_text(
                json.dumps(profil_info, indent=2, ensure_ascii=False), encoding="utf-8"
            )
            _publish_tree_atomically(staging, dest)

        self._progress.current_file = ""
        self._progress.status = "done"
        return result

    def _anonymize_docx(self, path: Path, profile: AnonymProfile) -> Tuple[bool, int]:
        """Anonymisiert ein Word-Dokument."""
        if not DOCX_AVAILABLE:
            return False, 0
        if (
            _has_unverified_package_content(path, self.trusted_media_hashes)
            and not self.allow_unverified_media
        ):
            return False, 0

        doc = Document(str(path))
        count = 0

        sorted_replacements = _collect_replacements(profile)

        # Wortgrenzen-Patterns vorab kompilieren (verhindert kaputte Fragmente
        # wie "Person026sbesonderheiten" statt "Wahrnehmungsbesonderheiten",
        # wenn ein NER-erkannter Einzelwort-Name zufaellig Praefix eines
        # laengeren deutschen Kompositums ist).
        compiled_replacements = [
            (re.compile(r'\b' + re.escape(old) + r'\b'), new)
            for old, new in sorted_replacements
        ]

        def replace_in_paragraphs(paragraphs):
            nonlocal count
            for paragraph in paragraphs:
                p_text = paragraph.text
                if not p_text:
                    continue

                # WICHTIG: Immer auf dem VOLLEN, zusammenhaengenden Absatztext
                # pruefen/ersetzen -- NIEMALS pro Run isoliert. Word splittet
                # ein einzelnes Wort haeufig auf mehrere interne Runs auf
                # (Formatierung, Autokorrektur, Bearbeitungshistorie); eine
                # Wortgrenzen-Pruefung auf Run-Ebene saehe dann faelschlich
                # eine saubere Grenze, obwohl der naechste Run das Wort nahtlos
                # fortsetzt (Fragment-Korruption trotz Wortgrenzen-Regex).
                # Fuer ein Datenschutz-Tool zaehlt Korrektheit mehr als der
                # Erhalt von Run-Formatierung.
                new_p_text = p_text
                changed = False
                for pattern, new in compiled_replacements:
                    new_p_text, n = pattern.subn(lambda m: new, new_p_text)
                    if n:
                        count += n
                        changed = True

                if changed and new_p_text != p_text:
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = new_p_text
                    else:
                        paragraph.add_run(new_p_text)

        # Paragraphen
        replace_in_paragraphs(doc.paragraphs)

        # Tabellen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

        # Header/Footer
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    replace_in_paragraphs(header.paragraphs)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    replace_in_paragraphs(footer.paragraphs)

        doc.core_properties.title, occurrences = _replace_mapped_text(
            doc.core_properties.title or "", sorted_replacements
        )
        count += occurrences
        doc.core_properties.subject, occurrences = _replace_mapped_text(
            doc.core_properties.subject or "", sorted_replacements
        )
        count += occurrences
        doc.core_properties.author, occurrences = _replace_mapped_text(
            doc.core_properties.author or "", sorted_replacements
        )
        count += occurrences
        doc.core_properties.keywords, occurrences = _replace_mapped_text(
            doc.core_properties.keywords or "", sorted_replacements
        )
        count += occurrences
        doc.core_properties.comments, occurrences = _replace_mapped_text(
            doc.core_properties.comments or "", sorted_replacements
        )
        count += occurrences

        doc.save(str(path))
        count += _sanitize_ooxml_package(path, sorted_replacements)
        _verify_ooxml_no_residuals(path, sorted_replacements)
        return True, count

    def _anonymize_text(self, path: Path, profile: AnonymProfile) -> Tuple[bool, int]:
        """Anonymisiert eine Textdatei."""
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")

        count = 0
        sorted_replacements = _collect_replacements(profile)

        for old, new in sorted_replacements:
            text, occurrences = _replace_word_boundary(text, old, new)
            count += occurrences

        path.write_text(text, encoding="utf-8")
        return True, count

    def _anonymize_doc(self, path: Path, profile: AnonymProfile) -> Tuple[bool, int]:
        """
        "Anonymisiert" ein altes Word-Binaerdokument (.doc).

        python-docx kann das binaere .doc-Format NICHT schreiben -- ohne
        diese Methode wird die Datei nur roh kopiert und bleibt zu 100%
        im Klartext (empirisch gefunden: Aktendeckblatt.doc landete
        unveraendert in data_ano/, Name/Diagnose/E-Mail vollstaendig lesbar).

        Workaround: Text extrahieren (antiword/LibreOffice), Ersetzungen
        anwenden, als GLEICHNAMIGE .txt-Datei speichern und das Original
        .doc loeschen. document_pipeline.py/DocumentCollector erkennt den
        Dokumenttyp anhand des Dateinamens (nicht der Endung), daher bleibt
        die Kategorisierung (z.B. "aktendeckblatt") beim Bundling erhalten.
        """
        text = _extract_legacy_doc_text(str(path))
        if not text:
            # Extraktion fehlgeschlagen (kein antiword/soffice verfuegbar o.ae.)
            # -- Datei NICHT unveraendert im "anonymisierten" Ordner belassen.
            path.unlink(missing_ok=True)
            return False, 0

        count = 0
        sorted_replacements = _collect_replacements(profile)

        for old, new in sorted_replacements:
            text, occurrences = _replace_word_boundary(text, old, new)
            count += occurrences

        txt_path = path.with_suffix(".txt")
        txt_path.write_text(text, encoding="utf-8")
        path.unlink(missing_ok=True)
        return True, count

    def _anonymize_excel(self, path: Path, profile: AnonymProfile) -> Tuple[bool, int]:
        """
        Anonymisiert eine Excel-Datei (.xlsx, .xls).

        Ersetzt sensible Begriffe in allen Zellen aller Tabellenblätter.
        """
        if not EXCEL_AVAILABLE:
            return False, 0
        if (
            _has_unverified_package_content(path, self.trusted_media_hashes)
            and not self.allow_unverified_media
        ):
            return False, 0

        try:
            wb = openpyxl.load_workbook(str(path), keep_links=False)
        except Exception:
            return False, 0

        count = 0
        sorted_replacements = _collect_replacements(profile)

        date_mappings = profile.mappings.get("dates", {})

        # Alle Tabellenblätter durchgehen
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            new_title, occurrences = _replace_mapped_text(sheet.title, sorted_replacements)
            if new_title != sheet.title:
                sheet.title = new_title[:31]
            count += occurrences

            for header_footer in (
                sheet.oddHeader, sheet.evenHeader, sheet.firstHeader,
                sheet.oddFooter, sheet.evenFooter, sheet.firstFooter,
            ):
                for part in (header_footer.left, header_footer.center, header_footer.right):
                    if isinstance(part.text, str):
                        part.text, occurrences = _replace_mapped_text(part.text, sorted_replacements)
                        count += occurrences
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        original = cell.value
                        new_value = original
                        for old, new in sorted_replacements:
                            new_value, occurrences = _replace_word_boundary(new_value, old, new)
                            count += occurrences
                        if new_value != original:
                            cell.value = new_value
                    elif isinstance(cell.value, (datetime, date)) and date_mappings:
                        # Excel speichert Datumszellen (z.B. Geburtsdatum in
                        # Zeitnachweis-Tabellen) oft als natives datetime/date-
                        # Objekt, NICHT als Text "TT.MM.JJJJ" -- der str()-Zweig
                        # oben griff hier nie, das echte Datum blieb also in
                        # jeder Datums-formatierten Zelle unanonymisiert.
                        cell_date_str = cell.value.strftime("%d.%m.%Y")
                        fake_str = date_mappings.get(cell_date_str)
                        if fake_str:
                            try:
                                fake_dt = datetime.strptime(fake_str, "%d.%m.%Y")
                                if isinstance(cell.value, datetime):
                                    cell.value = fake_dt.replace(
                                        hour=cell.value.hour,
                                        minute=cell.value.minute,
                                        second=cell.value.second,
                                    )
                                else:
                                    cell.value = fake_dt.date()
                                count += 1
                            except ValueError:
                                pass

                    if cell.comment is not None:
                        cell.comment.text, occurrences = _replace_mapped_text(
                            cell.comment.text or "", sorted_replacements
                        )
                        count += occurrences
                        cell.comment.author, occurrences = _replace_mapped_text(
                            cell.comment.author or "", sorted_replacements
                        )
                        count += occurrences

                    if cell.hyperlink is not None:
                        for attribute in ("target", "location", "tooltip", "display"):
                            value = getattr(cell.hyperlink, attribute, None)
                            if isinstance(value, str):
                                replaced, occurrences = _replace_mapped_text(value, sorted_replacements)
                                setattr(cell.hyperlink, attribute, replaced)
                                count += occurrences

        for attribute in (
            "title", "subject", "creator", "keywords", "description",
            "lastModifiedBy", "category", "contentStatus", "identifier",
            "language", "version",
        ):
            value = getattr(wb.properties, attribute, None)
            if isinstance(value, str):
                replaced, occurrences = _replace_mapped_text(value, sorted_replacements)
                setattr(wb.properties, attribute, replaced)
                count += occurrences

        wb.save(str(path))
        wb.close()
        try:
            count += _sanitize_ooxml_package(path, sorted_replacements)
            _verify_ooxml_no_residuals(path, sorted_replacements)
            verification = openpyxl.load_workbook(str(path), data_only=False, keep_links=False)
            verification.close()
        except Exception:
            return False, 0
        return True, count

    def _anonymize_filename(self, filepath: Path, profile: AnonymProfile) -> Path:
        """
        Anonymisiert einen Dateinamen, falls er sensible Begriffe enthält.

        Returns:
            Neuer Pfad (umbenannt) oder ursprünglicher Pfad (unverändert)
        """
        filename = filepath.stem
        suffix = filepath.suffix

        sorted_replacements = _collect_replacements(profile)

        new_filename = filename
        changed = False
        for old, new in sorted_replacements:
            if old in new_filename:
                new_filename = new_filename.replace(old, new)
                changed = True

        if changed:
            new_path = filepath.parent / f"{new_filename}{suffix}"
            filepath.rename(new_path)
            return new_path
        return filepath

    def _anonymize_pdf(self, path: Path, profile: AnonymProfile,
                       encrypt_password: Optional[str] = None) -> Tuple[bool, int]:
        """
        PDF-Anonymisierung via PyMuPDF-Schwärzung und optionale Verschlüsselung.

        Herkunfts-Komponentenreferenz (NICHT importiert — inline implementiert):
          - DokuZentrum RedactionDetector (Erkennungslogik)
          - PDFSchwaerzer Pro (Redact+Encrypt Pipeline)

        Pipeline:
          1. PDF öffnen (fitz)
          2. Für jede Seite: Sensitive Begriffe suchen und schwärzen
          3. Geschwärztes PDF speichern
          4. Optional: AES-256 verschlüsseln (pikepdf, R=6)

        Args:
            path: PDF-Datei
            profile: Anonymisierungsprofil mit Mappings
            encrypt_password: Optionales Passwort für PDF-Verschlüsselung
        """
        if not FITZ_AVAILABLE:
            return False, 0

        replacements = _collect_replacements(profile)
        sensitive_words = [old for old, _ in replacements if len(old) >= 2]

        count = 0
        temp_path = path.with_suffix(".tmp.pdf")

        try:
            doc = fitz.open(str(path))

            if not self.allow_unverified_media:
                for page in doc:
                    if page.get_images(full=True):
                        doc.close()
                        return False, 0

            for page in doc:
                for word in sensitive_words:
                    hits = page.search_for(word)
                    for rect in hits:
                        page.add_redact_annot(rect, fill=(0, 0, 0))
                        count += 1
                page.apply_redactions()

                annotations = list(page.annots() or [])
                for annotation in annotations:
                    page.delete_annot(annotation)
                for link in list(page.get_links() or []):
                    page.delete_link(link)
                if hasattr(page, "widgets") and hasattr(page, "delete_widget"):
                    for widget in list(page.widgets() or []):
                        page.delete_widget(widget)

            if hasattr(doc, "embfile_names") and hasattr(doc, "embfile_del"):
                for embedded_name in list(doc.embfile_names() or []):
                    doc.embfile_del(embedded_name)
            if hasattr(doc, "set_toc"):
                doc.set_toc([])
            if hasattr(doc, "set_page_labels"):
                doc.set_page_labels([])
            doc.set_metadata({})
            if hasattr(doc, "del_xml_metadata"):
                doc.del_xml_metadata()

            # Immer in Temp-Datei speichern (fitz verbietet non-incremental save zum Original)
            doc.save(str(temp_path), garbage=4, clean=True, deflate=True)
            doc.close()

            # Optional: AES-256 Verschlüsselung (pikepdf, R=6)
            if encrypt_password:
                if not PIKEPDF_AVAILABLE:
                    temp_path.unlink(missing_ok=True)
                    return False, 0
                try:
                    pdf = pikepdf.open(str(temp_path))
                    enc = pikepdf.Encryption(
                        owner=encrypt_password,
                        user=encrypt_password,
                        R=6
                    )
                    pdf.save(str(path), encryption=enc)
                    pdf.close()
                    temp_path.unlink()
                except Exception:
                    temp_path.unlink(missing_ok=True)
                    return False, 0
            else:
                if temp_path.exists():
                    shutil.move(str(temp_path), str(path))

            verification = fitz.open(str(path))
            if getattr(verification, "needs_pass", False):
                if not encrypt_password or not verification.authenticate(encrypt_password):
                    verification.close()
                    return False, 0
            residual_text = "\n".join(page.get_text() for page in verification)
            metadata_text = json.dumps(verification.metadata or {}, ensure_ascii=False)
            if _residual_originals(residual_text + "\n" + metadata_text, replacements):
                verification.close()
                return False, 0
            if not self.allow_unverified_media:
                for page in verification:
                    if page.get_images(full=True):
                        verification.close()
                        return False, 0
            if hasattr(verification, "embfile_names") and verification.embfile_names():
                verification.close()
                return False, 0
            if hasattr(verification, "get_toc") and verification.get_toc():
                verification.close()
                return False, 0
            if hasattr(verification, "get_page_labels") and verification.get_page_labels():
                verification.close()
                return False, 0
            verification.close()

            return True, count

        except Exception:
            if temp_path.exists():
                temp_path.unlink()
            return False, 0


# ═══════════════════════════════════════════════════════════════
# De-Anonymisierer
# ═══════════════════════════════════════════════════════════════

class DocumentDeanonymizer:
    """
    Stellt anonymisierte Dokumente wieder her.
    Ersetzt Tarnnamen durch echte Namen basierend auf dem Schlüssel.
    """

    def deanonymize_file(
        self,
        filepath: str,
        profile: AnonymProfile,
        trusted_template_path: Optional[str] = None,
    ) -> Tuple[bool, int]:
        """
        De-anonymisiert eine einzelne Datei (umgekehrte Mappings).

        trusted_template_path: siehe DocumentAnonymizer — erlaubt die
        Wiederherstellung von DOCX/XLSX-Ausgaben, deren eingebettete Medien
        byte-identisch aus einem vertrauenswuerdigen Vorlagen-Template stammen
        (sonst faellt die Deanonymisierung jedes Template-basierten Berichts
        auf den generellen unverifizierten-Medien-Stop zurueck).
        """
        # Umgekehrte Mappings erstellen
        reverse_profile = AnonymProfile(
            client_id=profile.client_id,
            tarnname="",
            fake_geburtsdatum="",
            mappings={}
        )

        for category, mapping in profile.mappings.items():
            reverse_profile.mappings[category] = {v: k for k, v in mapping.items()}

        # Anonymizer mit umgekehrten Mappings nutzen
        anon = DocumentAnonymizer(trusted_template_path=trusted_template_path)
        return anon.anonymize_file(filepath, reverse_profile)

    def deanonymize_folder(
        self,
        folder: str,
        schluessel_path: str,
        password: str,
        output_folder: str,
        client_id: str = None,
        trusted_template_path: Optional[str] = None,
    ) -> AnonymResult:
        """
        De-anonymisiert alle Dokumente in einem Ordner.

        Args:
            folder: Anonymisierter Ordner (z.B. klienten/K_0042/)
            schluessel_path: Pfad zur .schluessel.enc Datei (oder None wenn client_id gegeben)
            password: Passwort für den Schlüssel
            output_folder: Zielordner (z.B. _ready_for_export/Max_Mustermann/)
            trusted_template_path: siehe DocumentAnonymizer/deanonymize_file
            client_id: Klienten-ID — wenn angegeben, wird der lokale Schlüssel genutzt
        """
        # Schlüssel laden
        if client_id and not schluessel_path:
            schluessel_path = str(get_key_path(client_id))
        profile = decrypt_key_file(schluessel_path, password)

        result = AnonymResult()
        src = Path(folder)
        dest = Path(output_folder)
        _reject_symlink(src)
        if not src.is_dir():
            raise ValueError("Source folder must be an existing directory")
        source_root = src.resolve(strict=True)
        destination = dest.expanduser().resolve(strict=False)
        if _is_cloud_path(destination):
            raise ValueError("Restored plaintext must not be written to a cloud-synced path")
        if destination == source_root or _is_relative_to(destination, source_root):
            raise ValueError("Destination must be outside the source tree")
        _reject_symlink(dest, allow_missing_leaf=True)
        if dest.exists() and any(dest.iterdir()):
            raise ValueError("Destination must be absent or empty")

        # Dateien kopieren und de-anonymisieren
        deanon_suffixes = {".docx", ".txt", ".md"}
        copy_only_suffixes = {".pdf"}
        all_suffixes = deanon_suffixes | copy_only_suffixes

        files: List[Path] = []
        for root, directories, names in os.walk(source_root, followlinks=False):
            root_path = Path(root)
            for directory in list(directories):
                if _is_link_like(root_path / directory):
                    result.errors.append("link-like directory")
                    directories.remove(directory)
            for name in names:
                candidate = root_path / name
                if name.startswith("."):
                    continue
                if _is_link_like(candidate):
                    result.errors.append("link-like file")
                elif candidate.suffix.lower() in all_suffixes:
                    files.append(candidate)

        reverse_profile = AnonymProfile(
            client_id=profile.client_id,
            tarnname="",
            fake_geburtsdatum="",
            mappings={category: {value: key for key, value in mapping.items()}
                      for category, mapping in profile.mappings.items()},
        )

        with tempfile.TemporaryDirectory(prefix="deanonymizer_folder_") as staging_dir:
            staging = Path(staging_dir) / "publish"
            staging.mkdir()
            seen: set[Path] = set()
            for index, filepath in enumerate(files, start=1):
                result.processed_files += 1
                try:
                    rel = _anonymize_relative_path(filepath.relative_to(source_root), reverse_profile)
                    if rel in seen:
                        raise ValueError("restored path collision")
                    seen.add(rel)
                    staged_file = staging / rel
                    staged_file.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copyfile(filepath, staged_file)

                    if filepath.suffix.lower() in copy_only_suffixes:
                        continue
                    success, count = self.deanonymize_file(
                        str(staged_file), profile, trusted_template_path=trusted_template_path
                    )
                    if not success:
                        raise ValueError("de-anonymization failed")
                    result.anonymized_files += 1
                    result.replacements_total += count
                except Exception:
                    result.errors.append(f"file-{index}: de-anonymization failed")

            if result.errors:
                result.anonymized_files = 0
                result.replacements_total = 0
                result.errors.append("No restored output was published because validation failed")
                return result
            _publish_tree_atomically(staging, dest)

        return result


# ═══════════════════════════════════════════════════════════════
# CLI Interface
# ═══════════════════════════════════════════════════════════════

def main(argv: Optional[List[str]] = None) -> int:
    """Sicherer CLI-Einstiegspunkt ohne Geheimnisse in Prozessargumenten."""
    import argparse
    import getpass

    parser = argparse.ArgumentParser(prog="anonymizer", description="Anonymizer-Modul v0.2.1")
    subparsers = parser.add_subparsers(dest="command", required=True)
    subparsers.add_parser("self-test", aliases=["test"], help="lokalen Selbsttest ausführen")

    anonymize_parser = subparsers.add_parser("anonymize", help="Ordner anonymisieren")
    anonymize_parser.add_argument("source")
    anonymize_parser.add_argument("output")
    anonymize_parser.add_argument(
        "--allow-reduced-ner", action="store_true",
        help="bewusst ohne vollständige NER-Abdeckung arbeiten",
    )
    anonymize_parser.add_argument(
        "--trusted-template", dest="trusted_template", default=None,
        help=(
            "Pfad zu einer vertrauenswürdigen DOCX/XLSX-Vorlage; deren "
            "word/media|xl/media-Einträge dürfen byte-identisch übernommen "
            "werden (sonst ANONYMIZER_TRUSTED_TEMPLATE-Env)"
        ),
    )

    deanonymize_parser = subparsers.add_parser("deanonymize", help="Ordner lokal wiederherstellen")
    deanonymize_parser.add_argument("source")
    deanonymize_parser.add_argument("key")
    deanonymize_parser.add_argument("output")
    deanonymize_parser.add_argument(
        "--trusted-template", dest="trusted_template", default=None,
        help="siehe anonymize --trusted-template",
    )

    args = parser.parse_args(argv)

    if args.command in ("self-test", "test"):
        print("[TEST] Erstelle Testprofil...")
        anon = DocumentAnonymizer(require_ner=False)
        profile = anon.create_profile(
            real_name="Max Mustermann",
            geburtsdatum="15.03.2016",
            weitere_namen=["Dr. Meyer", "Frau Schmidt"],
            weitere_daten={"adresse": "Musterstr. 5, 79713 Bad Säckingen", "telefon": "07761/123456"}
        )
        print(f"  Client-ID: {profile.client_id}")
        print(f"  Tarnname : {profile.tarnname}")
        print(f"  Fake-Geb.: {profile.fake_geburtsdatum}")
        print("  Mappings:")
        for cat, mapping in profile.mappings.items():
            print(f"    {cat}:")
            for k, v in mapping.items():
                print(f"      {k} -> {v}")

        # Verschlüsselungstest
        if CRYPTO_AVAILABLE:
            import tempfile
            with tempfile.TemporaryDirectory() as td:
                test_path = Path(td) / "__test_schluessel.enc"
                encrypt_key_file(profile, str(test_path), "testpasswort123")
                print(f"\n  Schlüssel gespeichert: {test_path}")

                loaded = decrypt_key_file(str(test_path), "testpasswort123")
                print(f"  Entschlüsselt: {loaded.tarnname} (ID: {loaded.client_id})")
                assert loaded.tarnname == profile.tarnname
                assert loaded.mappings == profile.mappings
                print("  [OK] Verschlüsselungstest bestanden")
        else:
            print("\n  [WARN] cryptography nicht installiert, Verschlüsselungstest übersprungen")

        print("\n[OK] Test abgeschlossen")
        return 0

    if args.command == "anonymize":
        real_name = getpass.getpass("Echter Name (verborgene Eingabe): ")
        birth_date = getpass.getpass("Geburtsdatum TT.MM.JJJJ (verborgene Eingabe): ")
        password = getpass.getpass("Schlüsselpasswort (mindestens 12 Zeichen): ")
        anonymizer = DocumentAnonymizer(
            require_ner=not args.allow_reduced_ner,
            trusted_template_path=args.trusted_template,
        )
        scanned = anonymizer.scan_folder_for_sensitive_data(args.source)
        profile = anonymizer.create_profile(real_name, birth_date, scanned_data=scanned)
        result = anonymizer.anonymize_folder(
            args.source, profile, password, output_folder=args.output
        )
        print(
            f"Verarbeitet: {result.processed_files}; "
            f"veröffentlicht: {result.anonymized_files}; Fehler: {len(result.errors)}"
        )
        return 0 if not result.errors else 1

    if args.command == "deanonymize":
        password = getpass.getpass("Schlüsselpasswort: ")
        result = DocumentDeanonymizer().deanonymize_folder(
            args.source, args.key, password, args.output,
            trusted_template_path=args.trusted_template,
        )
        print(
            f"Verarbeitet: {result.processed_files}; "
            f"wiederhergestellt: {result.anonymized_files}; Fehler: {len(result.errors)}"
        )
        return 0 if not result.errors else 1

    return 2


if __name__ == "__main__":
    raise SystemExit(main())
