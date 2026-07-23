#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Defensive regression tests for the anonymizer privacy boundary."""

from __future__ import annotations

import json
import os
import hashlib
import tempfile
import unittest
import zipfile
from pathlib import Path
from types import SimpleNamespace
from unittest import mock

import anonymizer_modul.core as core
from anonymizer_modul.core import (
    AnonymProfile,
    CRYPTO_AVAILABLE,
    DOCX_AVAILABLE,
    EXCEL_AVAILABLE,
    FITZ_AVAILABLE,
    DocumentAnonymizer,
    decrypt_key_file,
    encrypt_key_file,
    get_key_path,
    main,
)


ORIGINAL = "Max Mustermann"
PSEUDONYM = "Elias Hartmann"


def profile() -> AnonymProfile:
    return AnonymProfile(
        client_id="K_ABC123",
        tarnname=PSEUDONYM,
        fake_geburtsdatum="12.04.2016",
        mappings={"names": {ORIGINAL: PSEUDONYM}},
    )


class FakeToken:
    """Minimal spaCy-Token-Double fuer die POS-/Lemma-basierte NER-Span-
    Validierung (nur die tatsaechlich genutzten Attribute: text, pos_,
    lemma_, idx, whitespace_/text_with_ws)."""

    def __init__(self, text, pos="PROPN", lemma=None, idx=0, ws=" "):
        self.text = text
        self.pos_ = pos
        self.lemma_ = lemma if lemma is not None else text
        self.idx = idx
        self.whitespace_ = ws

    @property
    def text_with_ws(self):
        return self.text + self.whitespace_


def fake_tokens(*specs):
    """Baut eine Liste FakeToken aus (text, pos[, lemma])-Tupeln mit
    fortlaufenden Zeichen-Offsets; das letzte Token bekommt keinen
    trailing Whitespace."""
    tokens = []
    cursor = 0
    for spec in specs:
        text, pos = spec[0], spec[1]
        lemma = spec[2] if len(spec) > 2 else text
        tokens.append(FakeToken(text, pos=pos, lemma=lemma, idx=cursor))
        cursor += len(text) + 1
    if tokens:
        tokens[-1].whitespace_ = ""
    return tokens


class FakeEnt:
    """Minimaler spaCy-Span-Double: iterierbar ueber FakeToken, mit label_
    fuer den PER/PERSON-Filter in detect_person_names_ner()."""

    def __init__(self, tokens, label="PER"):
        self._tokens = tokens
        self.label_ = label
        self.text = "".join(t.text_with_ws for t in tokens).rstrip()

    def __iter__(self):
        return iter(self._tokens)

    def __len__(self):
        return len(self._tokens)


class TestDetectionBoundary(unittest.TestCase):
    def test_ooxml_identity_attributes_enter_discovery_without_cross_value_matches(self):
        email = "Attribute.Person@Clinic.invalid"
        payload = (
            '<root xmlns:w="urn:fixture" w:val="' + email
            + '" descr="Case owner" technical="0006063" other="00034616" />'
        ).encode("utf-8")
        extracted = core._extract_xml_payload_text(payload)
        found = DocumentAnonymizer(require_ner=False).scan_text_for_sensitive_data(extracted)
        self.assertIn(email, found["emails"])
        self.assertEqual(found["phones"], [])

    def test_all_email_domains_are_sensitive(self):
        anonymizer = DocumentAnonymizer(require_ner=False)
        found = anonymizer.scan_text_for_sensitive_data(
            "alex.beispiel@schule.invalid und alex.beispiel@gmail.com"
        )
        self.assertEqual(
            found["emails"],
            ["alex.beispiel@schule.invalid", "alex.beispiel@gmail.com"],
        )

    def test_missing_ner_fails_closed_by_default(self):
        with mock.patch.object(core, "SPACY_AVAILABLE", False):
            with self.assertRaises(RuntimeError):
                DocumentAnonymizer().scan_text_for_sensitive_data("Max Mustermann")

            reduced = DocumentAnonymizer(require_ner=False)
            self.assertEqual(
                reduced.scan_text_for_sensitive_data("neutraler Text")["ner_person_names"],
                [],
            )

    def test_installed_spacy_without_a_configured_model_fails_closed(self):
        with (
            mock.patch.object(core, "SPACY_AVAILABLE", True),
            mock.patch.object(core, "_get_spacy_model", return_value=None),
        ):
            with self.assertRaises(RuntimeError):
                DocumentAnonymizer().scan_text_for_sensitive_data("Max Mustermann")

    def test_ambiguous_ner_volume_fails_closed(self):
        class FakeModel:
            max_length = 10_000

            def __call__(self, text):
                # idx bewusst ans Chunk-Ende gelegt (wie das vorherige
                # end_char=len(text)): chunk[idx:idx+1] ist dann "" (out of
                # bounds), damit der Wortgrenzen-Check hier nicht eingreift
                # -- dieser Test prueft ausschliesslich das Mengenlimit.
                # Zwei PROPN-Tokens pro Entitaet (Vorname+Nachname-Muster),
                # damit sie ueber das 0.2.3-Anker-Prinzip OHNE Lexikon-/
                # Titel-Anker als bestaetigt gelten (Mehrwort-Spans sind
                # immer verifiziert).
                entities = [
                    FakeEnt([
                        FakeToken(
                            f"Vorname{chr(65 + index // 26)}{chr(65 + index % 26)}",
                            pos="PROPN", idx=len(text), ws=" ",
                        ),
                        FakeToken(
                            f"Nachname{chr(65 + index // 26)}{chr(65 + index % 26)}",
                            pos="PROPN", idx=len(text), ws="",
                        ),
                    ])
                    for index in range(core._NER_MAX_NAMES_PER_CHUNK + 1)
                ]
                return SimpleNamespace(ents=entities)

        with (
            mock.patch.object(core, "SPACY_AVAILABLE", True),
            mock.patch.object(core, "NER_MODELS", ("fixture-model",)),
            mock.patch.object(core, "_get_spacy_model", return_value=FakeModel()),
            mock.patch.object(core, "_looks_like_person_name", return_value=True),
        ):
            with self.assertRaises(RuntimeError):
                DocumentAnonymizer().scan_text_for_sensitive_data("synthetic text")

    def test_extract_name_token_runs_shrinks_to_propn_subsequence(self):
        """Direkter Unit-Test der POS-basierten Span-Kuerzung ohne Modell:
        ein zu weit gefasster Span mit VERB/ADV/PRON-Rauschen um ein
        einzelnes PROPN-Token liefert nur dieses Token als Rest -- das
        strukturelle Gegenstueck zum realen "Grob bewegt Kim sich"-Fund."""
        tokens = fake_tokens(
            ("Grob", "ADV"), ("bewegt", "VERB"), ("Kim", "PROPN"), ("sich", "PRON"),
        )
        runs = core._extract_name_token_runs(FakeEnt(tokens))
        self.assertEqual(
            [core._tokens_to_name_text(run_tokens) for run_tokens, _preceding in runs],
            ["Kim"],
        )

    def test_tokens_pass_lemma_denylist_catches_inflected_forms(self):
        """Genitiv 'Landkreises' (Lemma 'Landkreis') wird ueber das Lemma
        erfasst, obwohl die Oberflaechenform nicht woertlich in der
        Denyliste steht."""
        self.assertFalse(
            core._tokens_pass_lemma_denylist(fake_tokens(("Landkreises", "NOUN", "Landkreis")))
        )
        self.assertTrue(
            core._tokens_pass_lemma_denylist(fake_tokens(("Kim", "PROPN", "Kim")))
        )

    def test_looks_like_person_name_rejects_generic_report_and_admin_nouns(self):
        """Direct filter-level regression for the observed foerderplaner
        Referenzlauf-Overblocking (2026-07-23): generic institutional/report
        nouns fused with or standing in for a proper name must not pass,
        while genuine synthetic names still do."""
        for candidate in (
            "Landkreis Lörrach", "Landkreis Loerrach", "Jugendamt Musterstadt",
            "Förderung", "Foerderung", "Zusage", "Ablauf",
        ):
            self.assertFalse(
                core._looks_like_person_name(candidate),
                f"{candidate!r} should not pass as a plausible person name",
            )
        for candidate in ("Kim", "Kim Beispiel", "Anna Muster", "Dr. Anna Muster"):
            self.assertTrue(
                core._looks_like_person_name(candidate),
                f"{candidate!r} should still pass as a plausible person name",
            )

    def test_run_has_anchor_single_word_needs_lexicon_or_title(self):
        """Direkter Unit-Test des 0.2.3-Anker-Prinzips (Operator-Design
        nach der RUN3-Regression: reines POS==PROPN kippt in der Praxis zu
        durchlaessig). Mehrwort-Spans sind immer verifiziert; Einzelwort-
        Spans brauchen einen Anker (Lexikon-Vorname ODER vorangehendes
        Titel-/Anrede-Token) -- sonst kein Anker, kein Ersatz."""
        # Mehrwort: immer Anker, auch ohne Lexikon-Eintrag ("Amara Diallo").
        self.assertTrue(core._run_has_anchor(fake_tokens(("Amara", "PROPN"), ("Diallo", "PROPN"))))
        # Einzelwort mit Lexikon-Vorname: Anker.
        self.assertTrue(core._run_has_anchor(fake_tokens(("Anna", "PROPN"))))
        # Einzelwort mit vorangehendem Titel-Token: Anker.
        self.assertTrue(core._run_has_anchor(
            fake_tokens(("Kim", "PROPN")), preceding_token=FakeToken("Dr.", pos="NOUN")
        ))
        self.assertTrue(core._run_has_anchor(
            fake_tokens(("Kim", "PROPN")), preceding_token=FakeToken("Frau", pos="NOUN")
        ))
        # Einzelwort ohne jeden Anker: kein Ersatz (Review-Kandidat).
        self.assertFalse(core._run_has_anchor(fake_tokens(("Kim", "PROPN"))))
        self.assertFalse(core._run_has_anchor(
            fake_tokens(("Umgang", "NOUN")), preceding_token=FakeToken("dem", pos="DET")
        ))

    def test_multiword_span_trims_generic_lemma_before_acceptance(self):
        """0.2.4-Haertung (RUN4-Nachbefund): reines PROPN-Tagging war auch
        bei Mehrwort-Spans durchlaessig -- "Landkreises Loerrach" wurde
        trotz PER-Label mit uebernommen, wenn beide Tokens (fehlerhaft)
        PROPN getaggt waren. Simuliert genau diesen Worst Case direkt auf
        Token-Ebene (unabhaengig von spaCys tatsaechlicher Kontext-Laune):
        das Gattungsbegriff-Token "Landkreises" (Lemma "Landkreis", auf der
        Denyliste) wird aus dem Span GEKUERZT, der Rest ("Loerrach") bleibt
        als Einzeltoken uebrig und braucht seinerseits einen Anker."""
        tokens = [
            FakeToken("Landkreises", pos="PROPN", lemma="Landkreis", idx=0),
            FakeToken("Loerrach", pos="PROPN", lemma="Loerrach", idx=12, ws=""),
        ]
        runs = core._extract_name_token_runs(FakeEnt(tokens))
        self.assertEqual(
            [core._tokens_to_name_text(run_tokens) for run_tokens, _preceding in runs],
            ["Loerrach"],
        )
        # Als Einzeltoken ohne Vornamen-/Titel-Anker: kein Ersatz.
        run_tokens, preceding = runs[0]
        self.assertFalse(core._run_has_anchor(run_tokens, preceding))

    def test_multiword_span_trims_substantivized_verb_before_acceptance(self):
        """0.2.4-Haertung: substantivierte Verben ("Anziehen") behalten ihr
        kleingeschriebenes Infinitiv-Lemma auch wenn sie (fehlerhaft) als
        PROPN getaggt werden -- werden trotz PROPN-Tag aus dem Span
        gekuerzt. "Beim Anziehen" darf so nie als Name durchgehen."""
        tokens = [
            FakeToken("Beim", pos="ADP", lemma="bei", idx=0),
            FakeToken("Anziehen", pos="PROPN", lemma="anziehen", idx=5, ws=""),
        ]
        runs = core._extract_name_token_runs(FakeEnt(tokens))
        self.assertEqual(runs, [])

    @unittest.skipUnless(
        core.SPACY_AVAILABLE and core._get_spacy_model("de_core_news_lg") is not None,
        "de_core_news_lg model required for the real NER regression",
    )
    def test_ner_run2_regression_no_failure_class_tokens_confirmed_or_review(self):
        """Regression mit den ECHTEN Saetzen aus dem foerderplaner-
        Referenzlauf RUN2/RUN3 (2026-07-23): weder in den bestaetigten noch
        in den Review-Kandidaten duerfen die bekannten Fehlklassen-Tokens
        auftauchen ("Grob", "Umgang", "Klettverschluesse", "Landkreises",
        "Einrichtungsangaben", "Foerderung"/"Zusage"/"Ablauf"). "Kim" darf
        je nach Satzkontext bestaetigt ODER nur Review-Kandidat sein (ohne
        Anker in einem isolierten Einzelsatz ist Review korrekt -- der
        eigentliche Klientenname kommt in der echten Pipeline ueber
        create_profile(real_name=...) unabhaengig von NER ins Profil)."""
        forbidden = {
            "grob", "umgang", "klettverschlüsse", "klettverschluesse",
            "landkreises", "einrichtungsangaben", "förderung", "foerderung",
            "zusage", "ablauf", "gegenständen", "gegenstaenden",
            "schwierigkeiten", "anziehen", "essen", "zähneputzen",
            "zaehneputzen", "begleitung", "handlauf",
        }
        texts = [
            "Grob bewegt Kim sich altersgemäß, beim Umgang mit kleinen "
            "Gegenständen mit den Fingern braucht Kim noch etwas "
            "Unterstützung.",
            "Beim Anziehen gelingt es Kim, Klettverschlüsse selbst zu öffnen.",
            "Beim Essen kommt Kim allein zurecht, beim Zähneputzen braucht "
            "Kim noch verbale Hinweise.",
            "Frau Lena Lindner bestätigt, dass die Bewilligung des "
            "fiktiven Landkreises Lörrach noch bis zum 31.08.2026 gilt.",
            "Beim Essen kommt Kim allein zurecht.",
            "Kim hat noch Schwierigkeiten bei Knöpfen und Reißverschlüssen.",
            "Alle Namen, Daten und Einrichtungsangaben in diesem Bericht sind fiktiv.",
            "Der Antrag wurde im Rahmen der Bewilligung des fiktiven "
            "Landkreises Lörrach genehmigt.",
            "Aus fachlicher Sicht sollte die Förderung von Kim weitergehen.",
        ]
        for text in texts:
            with self.subTest(text=text):
                confirmed, review_only = core.detect_person_names_ner(text)
                all_hits = {n.lower() for n in confirmed + review_only}
                self.assertFalse(
                    all_hits & forbidden,
                    f"Fehlklassen-Token in {text!r}: {all_hits & forbidden}",
                )

    @unittest.skipUnless(
        core.SPACY_AVAILABLE and core._get_spacy_model("de_core_news_lg") is not None,
        "de_core_news_lg model required for the real NER regression",
    )
    def test_ner_run2_regression_positive_names_still_detected(self):
        """Echte synthetische Namen -- inklusive eines Namens OHNE
        Lexikon-Eintrag ('Amara Diallo') -- muessen nach dem strukturellen
        POS-/Anker-Fix weiterhin (bestaetigt ODER als Review-Kandidat)
        auftauchen, niemals komplett verschwinden."""
        cases = [
            (
                "Kim Beispiel wurde vorgestellt. Dr. Anna Muster war "
                "ebenfalls anwesend.",
                {"Kim", "Anna Muster"},
            ),
            ("Amara Diallo wurde vorgestellt.", {"Amara Diallo"}),
        ]
        for text, expected in cases:
            with self.subTest(text=text):
                confirmed, review_only = core.detect_person_names_ner(text)
                self.assertEqual(set(confirmed) | set(review_only), expected)
        # "Amara Diallo" ist ein Mehrwort-Span ohne Lexikon-Eintrag -- muss
        # ueber das Anker-Prinzip BESTAETIGT (nicht nur Review) sein.
        confirmed, _ = core.detect_person_names_ner("Amara Diallo wurde vorgestellt.")
        self.assertIn("Amara Diallo", confirmed)
        # 0.2.4: Bindestrich-Doppelname darf durch die neue Mehrwort-
        # Haertung NICHT zerkuerzt werden (kein Gattungsbegriff, kein
        # substantiviertes Verb -- beide Teile sind echte PROPN-Nachnamen).
        confirmed, _ = core.detect_person_names_ner("Anna Muster-Bergmann wurde vorgestellt.")
        self.assertIn("Anna Muster-Bergmann", confirmed)

    @unittest.skipUnless(
        core.SPACY_AVAILABLE and core._get_spacy_model("de_core_news_lg") is not None,
        "de_core_news_lg model required for the real end-to-end DoD check",
    )
    def test_ner_run2_real_akte_end_to_end_no_corruption(self):
        """Pflicht-DoD (Operator, nach RUN3-Regression): DocumentAnonymizer
        gegen die ECHTE synthetische RUN2-Akte
        (C:\\_Local_DEV\\foerderplaner_referenz\\ARCHIV_2026-07-23_RUN2\\quelle\\,
        synthetisch) -- die tatsaechlich anonymisierte Ausgabe darf KEINE
        der bekannten Fehlklassen enthalten und muss weiterhin lesbaren,
        unkorrumpierten Fliesstext liefern. Wird uebersprungen, wenn der
        Referenzordner auf diesem Host nicht vorhanden ist."""
        source = Path(r"C:\_Local_DEV\foerderplaner_referenz\ARCHIV_2026-07-23_RUN2\quelle")
        if not source.is_dir():
            self.skipTest("RUN2-Referenzakte auf diesem Host nicht vorhanden")

        anon = DocumentAnonymizer(require_ner=True)
        found = anon.scan_folder_for_sensitive_data(str(source))
        all_scan_hits = {n.lower() for n in found["ner_person_names"] + found["ner_review_only"]}
        forbidden = {
            "grob", "umgang", "klettverschlüsse", "klettverschluesse",
            "landkreises", "einrichtungsangaben", "förderung", "foerderung",
            "begleitung", "handlauf", "zähneputzen", "zaehneputzen",
        }
        self.assertFalse(all_scan_hits & forbidden, f"Fehlklassen im Scan: {all_scan_hits & forbidden}")

        profile = anon.create_profile(
            real_name="Kim Beispiel", geburtsdatum="01.02.2018", scanned_data=found
        )
        mapped = set(profile.mappings["names"].keys())
        self.assertIn("Kim", mapped)
        self.assertIn("Kim Beispiel", mapped)
        self.assertFalse({m.lower() for m in mapped} & forbidden)

        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "out"
            result = anon.anonymize_folder(
                str(source), profile, "ein-langes-lokales-testpasswort",
                output_folder=str(output),
            )
            self.assertEqual(result.errors, [])
            self.assertEqual(result.anonymized_files, 3)
            published_text = "\n".join(
                f.read_text(encoding="utf-8") for f in output.rglob("*.txt")
            ) + "\n".join(f.read_text(encoding="utf-8") for f in output.rglob("*.md"))
            lowered = published_text.lower()
            # Diese Fragmente enthalten kein "Kim" und muessen daher
            # UNVERAENDERT (nicht korrumpiert) im Output erhalten bleiben --
            # das ist der eigentliche RUN3-Regressionsbefund: 0.2.2 ersetzte
            # genau diese generischen Woerter faelschlich durch Fake-Namen.
            for intact_fragment in (
                "grob bewegt", "beim umgang", "klettverschlüsse",
                "landkreises lörrach", "einrichtungsangaben", "die förderung",
                "begleitung", "handlauf", "zähneputzen",
            ):
                self.assertIn(
                    intact_fragment, lowered,
                    f"Fragment {intact_fragment!r} fehlt/korrumpiert im Output",
                )
            # Der eigentliche Klientenname MUSS dagegen vollstaendig ersetzt sein.
            self.assertNotIn("kim", lowered)

    def test_mixed_case_email_is_replaced_without_lowercasing_the_source(self):
        email = "Alice.Example@Unknown.com"
        anonymizer = DocumentAnonymizer(require_ner=False)
        scanned = anonymizer.scan_text_for_sensitive_data(email)
        self.assertEqual(scanned["emails"], [email])
        created = anonymizer.create_profile(
            ORIGINAL,
            "15.03.2016",
            scanned_data=scanned,
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "case.txt"
            path.write_text(email, encoding="utf-8")
            success, count = anonymizer.anonymize_file(str(path), created)
            self.assertTrue(success)
            self.assertEqual(count, 1)
            self.assertNotIn(email, path.read_text(encoding="utf-8"))

    def test_detected_names_and_institutions_enter_profile(self):
        anonymizer = DocumentAnonymizer(require_ner=False)
        created = anonymizer.create_profile(
            ORIGINAL,
            "15.03.2016",
            scanned_data={
                "table_row_names": ["Erika Beispiel"],
                "ner_person_names": ["Noah Beispiel"],
                "institutions": ["Beispielschule"],
            },
        )
        self.assertIn("Erika Beispiel", created.mappings["names"])
        self.assertIn("Noah Beispiel", created.mappings["names"])
        self.assertIn("Beispielschule", created.mappings["institutions"])


class TestKeyBoundary(unittest.TestCase):
    def test_client_id_cannot_escape_key_directory(self):
        with tempfile.TemporaryDirectory() as temp_dir, mock.patch.dict(
            os.environ, {"ANONYMIZER_KEYS_DIR": temp_dir}
        ):
            with self.assertRaises(ValueError):
                get_key_path("../outside")
            self.assertEqual(get_key_path("K_ABC123").parent, Path(temp_dir).resolve())

    def test_cloud_key_override_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            cloud_path = Path(temp_dir) / "OneDrive-fixture" / "keys"
            with mock.patch.dict(os.environ, {"ANONYMIZER_KEYS_DIR": str(cloud_path)}):
                with self.assertRaises(ValueError):
                    get_key_path("K_ABC123")

    @unittest.skipUnless(CRYPTO_AVAILABLE, "cryptography required")
    def test_short_password_and_cloud_output_are_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            local_output = Path(temp_dir) / "key.enc"
            with self.assertRaises(ValueError):
                encrypt_key_file(profile(), str(local_output), "short")
            cloud_output = Path(temp_dir) / "OneDrive-fixture" / "key.enc"
            with self.assertRaises(ValueError):
                encrypt_key_file(profile(), str(cloud_output), "TestPasswort_2026!")

    @unittest.skipUnless(CRYPTO_AVAILABLE, "cryptography required")
    def test_oversized_key_blob_is_rejected_before_reading_payload(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "oversized.enc"
            with path.open("wb") as stream:
                stream.seek(10 * 1024 * 1024)
                stream.write(b"x")
            with self.assertRaises(ValueError):
                decrypt_key_file(str(path), "TestPasswort_2026!")


class TestTransactionalFolderBoundary(unittest.TestCase):
    def test_single_file_with_sensitive_filename_fails_closed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / f"Bericht {ORIGINAL}.txt"
            path.write_text(ORIGINAL, encoding="utf-8")
            success, count = DocumentAnonymizer(require_ner=False).anonymize_file(
                str(path), profile()
            )
            self.assertFalse(success)
            self.assertEqual(count, 0)
            self.assertEqual(path.read_text(encoding="utf-8"), ORIGINAL)

    def test_unsupported_file_never_publishes_output(self):
        with tempfile.TemporaryDirectory() as temp_dir, mock.patch.dict(
            os.environ, {"ANONYMIZER_KEYS_DIR": str(Path(temp_dir) / "keys")}
        ):
            source = Path(temp_dir) / "source"
            output = Path(temp_dir) / "output"
            source.mkdir()
            (source / "case.csv").write_text(ORIGINAL, encoding="utf-8")
            result = DocumentAnonymizer(require_ner=False).anonymize_folder(
                str(source), profile(), "TestPasswort_2026!", str(output)
            )
            self.assertTrue(result.errors)
            self.assertFalse(output.exists())
            self.assertNotIn("case.csv", "\n".join(result.errors))

    def test_format_failure_never_publishes_output(self):
        for suffix in (".docx", ".xlsx", ".pdf"):
            with self.subTest(suffix=suffix), tempfile.TemporaryDirectory() as temp_dir:
                source = Path(temp_dir) / "source"
                output = Path(temp_dir) / "output"
                source.mkdir()
                (source / f"case{suffix}").write_text(ORIGINAL, encoding="utf-8")
                with mock.patch.dict(
                    os.environ, {"ANONYMIZER_KEYS_DIR": str(Path(temp_dir) / "keys")}
                ):
                    result = DocumentAnonymizer(require_ner=False).anonymize_folder(
                        str(source), profile(), "TestPasswort_2026!", str(output)
                    )
                self.assertTrue(result.errors)
                self.assertFalse(output.exists())

    def test_source_destination_aliases_are_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "source"
            source.mkdir()
            (source / "case.txt").write_text(ORIGINAL, encoding="utf-8")
            anonymizer = DocumentAnonymizer(require_ner=False)
            with self.assertRaises(ValueError):
                anonymizer.anonymize_folder(
                    str(source), profile(), "TestPasswort_2026!", str(source)
                )
            with self.assertRaises(ValueError):
                anonymizer.anonymize_folder(
                    str(source), profile(), "TestPasswort_2026!", str(source / "nested")
                )
            with self.assertRaises(ValueError):
                anonymizer.anonymize_folder(
                    str(source), profile(), "TestPasswort_2026!",
                    str(Path(temp_dir) / ORIGINAL),
                )

    def test_source_symlink_is_not_published(self):
        with tempfile.TemporaryDirectory() as temp_dir, mock.patch.dict(
            os.environ, {"ANONYMIZER_KEYS_DIR": str(Path(temp_dir) / "keys")}
        ):
            source = Path(temp_dir) / "source"
            source.mkdir()
            outside = Path(temp_dir) / "outside.txt"
            outside.write_text(ORIGINAL, encoding="utf-8")
            link = source / "linked.txt"
            try:
                link.symlink_to(outside)
            except OSError as error:
                self.skipTest(f"symlinks unavailable: {error}")
            output = Path(temp_dir) / "output"
            result = DocumentAnonymizer(require_ner=False).anonymize_folder(
                str(source), profile(), "TestPasswort_2026!", str(output)
            )
            self.assertTrue(result.errors)
            self.assertFalse(output.exists())

    def test_destination_symlink_is_rejected(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "source"
            real_destination = Path(temp_dir) / "real-destination"
            source.mkdir()
            real_destination.mkdir()
            (source / "case.txt").write_text(ORIGINAL, encoding="utf-8")
            destination = Path(temp_dir) / "destination"
            try:
                destination.symlink_to(real_destination, target_is_directory=True)
            except OSError as error:
                self.skipTest(f"symlinks unavailable: {error}")
            with self.assertRaises(ValueError):
                DocumentAnonymizer(require_ner=False).anonymize_folder(
                    str(source), profile(), "TestPasswort_2026!", str(destination)
                )

    def test_windows_reparse_point_is_treated_as_link_like(self):
        metadata = SimpleNamespace(
            st_mode=core.stat.S_IFDIR,
            st_file_attributes=getattr(core.stat, "FILE_ATTRIBUTE_REPARSE_POINT", 0x400),
        )
        with mock.patch.object(core.os, "lstat", return_value=metadata):
            self.assertTrue(core._is_link_like(Path("junction-fixture")))

    def test_atomic_publication_removes_partial_sibling_on_copy_failure(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            staging = root / "staging"
            destination = root / "published"
            staging.mkdir()
            (staging / "verified.txt").write_text("verified", encoding="utf-8")

            def fail_after_partial_copy(source, candidate):
                candidate = Path(candidate)
                candidate.mkdir(parents=True)
                (candidate / "partial.txt").write_text("partial", encoding="utf-8")
                raise OSError("synthetic copy failure")

            with (
                mock.patch.object(core.shutil, "copytree", side_effect=fail_after_partial_copy),
                self.assertRaises(OSError),
            ):
                core._publish_tree_atomically(staging, destination)

            self.assertFalse(destination.exists())
            self.assertEqual(list(root.glob(".anonymizer-publish-*")), [])

    @unittest.skipUnless(CRYPTO_AVAILABLE, "cryptography required")
    def test_directory_names_are_anonymized_and_profile_hides_key_path(self):
        with tempfile.TemporaryDirectory() as temp_dir, mock.patch.dict(
            os.environ, {"ANONYMIZER_KEYS_DIR": str(Path(temp_dir) / "keys")}
        ):
            source = Path(temp_dir) / "source"
            nested = source / ORIGINAL
            nested.mkdir(parents=True)
            (nested / f"Bericht {ORIGINAL}.txt").write_text(
                f"Fallakte {ORIGINAL}", encoding="utf-8"
            )
            output = Path(temp_dir) / "output"
            result = DocumentAnonymizer(require_ner=False).anonymize_folder(
                str(source), profile(), "TestPasswort_2026!", str(output)
            )
            self.assertFalse(result.errors)
            paths = [str(path.relative_to(output)) for path in output.rglob("*")]
            self.assertFalse(any(ORIGINAL in path for path in paths))
            self.assertTrue(any(PSEUDONYM in path for path in paths))
            info = json.loads((output / ".profil.json").read_text(encoding="utf-8"))
            self.assertNotIn("key_location", info)


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx required")
class TestDocxBoundary(unittest.TestCase):
    def test_comments_and_properties_do_not_retain_original(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "case.docx"
            document = Document()
            paragraph = document.add_paragraph(ORIGINAL)
            document.add_comment(paragraph.runs[0], text=ORIGINAL, author=ORIGINAL)
            document.core_properties.subject = ORIGINAL
            document.save(path)

            success, count = DocumentAnonymizer(require_ner=False).anonymize_file(
                str(path), profile()
            )
            self.assertTrue(success)
            self.assertGreaterEqual(count, 3)
            with zipfile.ZipFile(path) as archive:
                xml = b"\n".join(
                    archive.read(info)
                    for info in archive.infolist()
                    if info.filename.endswith((".xml", ".rels"))
                ).decode("utf-8", errors="replace")
            self.assertNotIn(ORIGINAL, xml)

    @unittest.skipUnless(CRYPTO_AVAILABLE, "cryptography required")
    def test_scan_profile_publish_covers_secondary_package_surfaces(self):
        from docx import Document

        email = "Case.Person@Clinic.invalid"
        with tempfile.TemporaryDirectory() as temp_dir, mock.patch.dict(
            os.environ, {"ANONYMIZER_KEYS_DIR": str(Path(temp_dir) / "keys")}
        ):
            source = Path(temp_dir) / "source"
            output = Path(temp_dir) / "output"
            source.mkdir()
            path = source / "case.docx"
            document = Document()
            paragraph = document.add_paragraph("Neutral content")
            document.add_comment(paragraph.runs[0], text=email, author="Fixture")
            document.sections[0].header.paragraphs[0].text = email
            document.core_properties.subject = email
            document.save(path)

            anonymizer = DocumentAnonymizer(require_ner=False)
            scanned = anonymizer.scan_folder_for_sensitive_data(str(source))
            self.assertIn(email, scanned["emails"])
            created = anonymizer.create_profile(
                ORIGINAL,
                "15.03.2016",
                scanned_data=scanned,
            )
            result = anonymizer.anonymize_folder(
                str(source), created, "TestPasswort_2026!", str(output)
            )
            self.assertFalse(result.errors)
            with zipfile.ZipFile(output / "case.docx") as archive:
                xml = b"\n".join(
                    archive.read(info)
                    for info in archive.infolist()
                    if info.filename.endswith((".xml", ".rels"))
                ).decode("utf-8", errors="replace")
            self.assertNotIn(email.casefold(), xml.casefold())

    def test_unverified_media_fails_closed_and_preserves_original(self):
        from docx import Document
        from docx.shared import Inches
        import base64

        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "media.docx"
            image = Path(temp_dir) / "pixel.png"
            image.write_bytes(png)
            document = Document()
            document.add_paragraph(ORIGINAL)
            document.add_picture(str(image), width=Inches(0.1))
            document.save(path)
            before = hashlib.sha256(path.read_bytes()).hexdigest()
            success, count = DocumentAnonymizer(require_ner=False).anonymize_file(
                str(path), profile()
            )
            self.assertFalse(success)
            self.assertEqual(count, 0)
            self.assertEqual(hashlib.sha256(path.read_bytes()).hexdigest(), before)

    def test_trusted_template_media_hash_match_publishes(self):
        """Media byte-identical to the trusted template's media passes."""
        from docx import Document
        from docx.shared import Inches
        import base64

        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir = Path(temp_dir)
            image = temp_dir / "pixel.png"
            image.write_bytes(png)

            template_path = temp_dir / "template.docx"
            template = Document()
            template.add_paragraph("Vorlage")
            template.add_picture(str(image), width=Inches(0.1))
            template.save(template_path)

            document_path = temp_dir / "report.docx"
            document = Document()
            document.add_paragraph(ORIGINAL)
            document.add_picture(str(image), width=Inches(0.1))
            document.save(document_path)

            success, count = DocumentAnonymizer(
                require_ner=False, trusted_template_path=str(template_path)
            ).anonymize_file(str(document_path), profile())
            self.assertTrue(success)
            self.assertGreaterEqual(count, 1)

    def test_trusted_template_media_hash_mismatch_still_blocks(self):
        """A foreign image beside a template-matching one still fails closed."""
        from docx import Document
        from docx.shared import Inches
        import base64

        png_template = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        png_foreign = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir = Path(temp_dir)
            template_image = temp_dir / "template_pixel.png"
            template_image.write_bytes(png_template)
            foreign_image = temp_dir / "foreign_pixel.png"
            foreign_image.write_bytes(png_foreign)

            template_path = temp_dir / "template.docx"
            template = Document()
            template.add_paragraph("Vorlage")
            template.add_picture(str(template_image), width=Inches(0.1))
            template.save(template_path)

            document_path = temp_dir / "report.docx"
            document = Document()
            document.add_paragraph(ORIGINAL)
            document.add_picture(str(template_image), width=Inches(0.1))
            document.add_picture(str(foreign_image), width=Inches(0.1))
            document.save(document_path)
            before = hashlib.sha256(document_path.read_bytes()).hexdigest()

            success, count = DocumentAnonymizer(
                require_ner=False, trusted_template_path=str(template_path)
            ).anonymize_file(str(document_path), profile())
            self.assertFalse(success)
            self.assertEqual(count, 0)
            self.assertEqual(hashlib.sha256(document_path.read_bytes()).hexdigest(), before)


@unittest.skipUnless(EXCEL_AVAILABLE, "openpyxl required")
class TestXlsxBoundary(unittest.TestCase):
    def test_comments_metadata_and_hidden_sheet_are_sanitized(self):
        import openpyxl
        from openpyxl.comments import Comment

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "case.xlsx"
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet["A1"] = ORIGINAL
            sheet["A1"].comment = Comment(ORIGINAL, ORIGINAL)
            workbook.properties.subject = ORIGINAL
            hidden = workbook.create_sheet("Hidden")
            hidden.sheet_state = "hidden"
            hidden["A1"] = ORIGINAL
            workbook.save(path)
            workbook.close()

            success, count = DocumentAnonymizer(require_ner=False).anonymize_file(
                str(path), profile()
            )
            self.assertTrue(success)
            self.assertGreaterEqual(count, 4)
            reopened = openpyxl.load_workbook(path, data_only=False, keep_links=False)
            self.assertNotIn(ORIGINAL, reopened.active["A1"].value)
            self.assertNotIn(ORIGINAL, reopened.active["A1"].comment.text)
            self.assertNotIn(ORIGINAL, reopened.properties.subject or "")
            self.assertNotIn(ORIGINAL, reopened["Hidden"]["A1"].value)
            reopened.close()

    @unittest.skipUnless(CRYPTO_AVAILABLE, "cryptography required")
    def test_scan_profile_publish_covers_secondary_package_surfaces(self):
        import openpyxl
        from openpyxl.comments import Comment

        email = "Case.Person@Clinic.invalid"
        with tempfile.TemporaryDirectory() as temp_dir, mock.patch.dict(
            os.environ, {"ANONYMIZER_KEYS_DIR": str(Path(temp_dir) / "keys")}
        ):
            source = Path(temp_dir) / "source"
            output = Path(temp_dir) / "output"
            source.mkdir()
            path = source / "case.xlsx"
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet["A1"] = "Neutral content"
            sheet["A1"].comment = Comment(email, "Fixture")
            sheet.oddHeader.center.text = email
            workbook.properties.subject = email
            workbook.save(path)
            workbook.close()

            anonymizer = DocumentAnonymizer(require_ner=False)
            scanned = anonymizer.scan_folder_for_sensitive_data(str(source))
            self.assertIn(email, scanned["emails"])
            created = anonymizer.create_profile(
                ORIGINAL,
                "15.03.2016",
                scanned_data=scanned,
            )
            result = anonymizer.anonymize_folder(
                str(source), created, "TestPasswort_2026!", str(output)
            )
            self.assertFalse(result.errors)
            with zipfile.ZipFile(output / "case.xlsx") as archive:
                xml = b"\n".join(
                    archive.read(info)
                    for info in archive.infolist()
                    if info.filename.endswith((".xml", ".rels"))
                ).decode("utf-8", errors="replace")
            self.assertNotIn(email.casefold(), xml.casefold())

    @unittest.skipUnless(CRYPTO_AVAILABLE, "cryptography required")
    def test_scan_profile_publish_covers_chart_xml(self):
        import openpyxl
        from openpyxl.chart import BarChart, Reference

        email = "Chart.Person@Clinic.invalid"
        with tempfile.TemporaryDirectory() as temp_dir, mock.patch.dict(
            os.environ, {"ANONYMIZER_KEYS_DIR": str(Path(temp_dir) / "keys")}
        ):
            source = Path(temp_dir) / "source"
            output = Path(temp_dir) / "output"
            source.mkdir()
            path = source / "chart.xlsx"
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.append(["Value"])
            sheet.append([1])
            chart = BarChart()
            chart.title = email
            chart.add_data(
                Reference(sheet, min_col=1, min_row=1, max_row=2),
                titles_from_data=True,
            )
            sheet.add_chart(chart, "C3")
            workbook.save(path)
            workbook.close()

            anonymizer = DocumentAnonymizer(require_ner=False)
            scanned = anonymizer.scan_folder_for_sensitive_data(str(source))
            self.assertIn(email, scanned["emails"])
            created = anonymizer.create_profile(
                ORIGINAL,
                "15.03.2016",
                scanned_data=scanned,
            )
            result = anonymizer.anonymize_folder(
                str(source), created, "TestPasswort_2026!", str(output)
            )
            self.assertFalse(result.errors)
            with zipfile.ZipFile(output / "chart.xlsx") as archive:
                xml = b"\n".join(
                    archive.read(info)
                    for info in archive.infolist()
                    if info.filename.endswith((".xml", ".rels"))
                ).decode("utf-8", errors="replace")
            self.assertNotIn(email.casefold(), xml.casefold())

    def test_unverified_media_member_fails_closed(self):
        import openpyxl

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "media.xlsx"
            workbook = openpyxl.Workbook()
            workbook.active["A1"] = ORIGINAL
            workbook.save(path)
            workbook.close()
            with zipfile.ZipFile(path, "a") as archive:
                archive.writestr("xl/media/placeholder.txt", ORIGINAL)
            before = hashlib.sha256(path.read_bytes()).hexdigest()
            success, count = DocumentAnonymizer(require_ner=False).anonymize_file(
                str(path), profile()
            )
            self.assertFalse(success)
            self.assertEqual(count, 0)
            self.assertEqual(hashlib.sha256(path.read_bytes()).hexdigest(), before)


@unittest.skipUnless(FITZ_AVAILABLE, "PyMuPDF required")
class TestPdfBoundary(unittest.TestCase):
    def test_text_metadata_annotations_and_attachments_are_removed(self):
        import fitz

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "case.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), ORIGINAL)
            page.add_text_annot((80, 80), ORIGINAL)
            document.set_toc([[1, ORIGINAL, 1]])
            document.set_metadata({"title": ORIGINAL})
            if hasattr(document, "embfile_add"):
                document.embfile_add("note.txt", ORIGINAL.encode("utf-8"))
            document.save(path)
            document.close()

            success, count = DocumentAnonymizer(require_ner=False).anonymize_file(
                str(path), profile()
            )
            self.assertTrue(success)
            self.assertGreater(count, 0)
            reopened = fitz.open(path)
            text = "\n".join(page.get_text() for page in reopened)
            self.assertNotIn(ORIGINAL, text)
            self.assertNotIn(ORIGINAL, json.dumps(reopened.metadata or {}))
            self.assertEqual(list(reopened[0].annots() or []), [])
            self.assertEqual(reopened.get_toc(), [])
            if hasattr(reopened, "embfile_names"):
                self.assertEqual(reopened.embfile_names(), [])
            reopened.close()

    def test_image_content_fails_closed_and_preserves_original(self):
        import base64
        import fitz

        png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "image.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_image(fitz.Rect(10, 10, 20, 20), stream=png)
            document.save(path)
            document.close()
            before = hashlib.sha256(path.read_bytes()).hexdigest()
            success, count = DocumentAnonymizer(require_ner=False).anonymize_file(
                str(path), profile()
            )
            self.assertFalse(success)
            self.assertEqual(count, 0)
            self.assertEqual(hashlib.sha256(path.read_bytes()).hexdigest(), before)

    def test_requested_pdf_encryption_never_falls_back_to_plaintext(self):
        import fitz

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "encrypt.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), ORIGINAL)
            document.save(path)
            document.close()
            before = hashlib.sha256(path.read_bytes()).hexdigest()
            anonymizer = DocumentAnonymizer(require_ner=False)
            with mock.patch.object(core, "PIKEPDF_AVAILABLE", False):
                success, count = anonymizer._anonymize_pdf(
                    path, profile(), encrypt_password="TestPasswort_2026!"
                )
            self.assertFalse(success)
            self.assertEqual(count, 0)
            self.assertEqual(hashlib.sha256(path.read_bytes()).hexdigest(), before)


class TestLegacyDocBoundary(unittest.TestCase):
    def test_legacy_doc_size_limit_precedes_external_conversion(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "large.doc"
            with path.open("wb") as stream:
                stream.seek(25 * 1024 * 1024)
                stream.write(b"x")
            with self.assertRaises(ValueError):
                core._extract_legacy_doc_text(str(path))


class TestCliBoundary(unittest.TestCase):
    def test_self_test_command_is_real(self):
        self.assertEqual(main(["self-test"]), 0)

    def test_unknown_command_is_nonzero(self):
        with self.assertRaises(SystemExit) as caught:
            main(["not-a-command"])
        self.assertEqual(caught.exception.code, 2)


if __name__ == "__main__":
    unittest.main(verbosity=2)
